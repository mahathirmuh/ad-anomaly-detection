#!/usr/bin/env python3
"""
Generate Project Report DOCX v2 - CURRENT STATE
AD Anomaly Detection - Graph-Based Knowledge System
Mencerminkan kondisi project saat ini (semua phase selesai)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
import json
import io
import os

os.makedirs('output', exist_ok=True)

# ── ACTUAL PIPELINE RESULTS ──────────────────────────────────────────
with open('output/anomaly_statistics.json') as f:
    STATS = json.load(f)

TOTAL_USERS = STATS['total_users']
ANOMALY_DIST = STATS['anomaly_distribution']
SCORE_STATS = STATS['anomaly_score_stats']
ENSEMBLE_VOTING = STATS['ensemble_voting']

TOP_ANOMALIES = [
    ('mti.admin',           0.6800, 'HIGH',   '7/7', '3/3', 'Rasio login gagal'),
    ('mti.sysadmin',        0.4276, 'MEDIUM', '7/7', '3/3', 'Banyak admin action'),
    ('andre.saputra',       0.3870, 'LOW',    '7/7', '2/3', 'Banyak admin action'),
    ('peggy.putra',         0.3834, 'LOW',    '7/7', '2/3', 'Banyak admin action'),
    ('mahathir.muhammad',   0.3763, 'LOW',    '6/7', '2/3', 'Rasio login gagal'),
    ('eris.rismansyah',     0.3727, 'LOW',    '6/7', '2/3', 'Sering lockout'),
    ('ricky.rediansyah',    0.3562, 'LOW',    '1/7', '2/3', 'Pelanggaran rule'),
    ('aini.rosyidah',       0.3562, 'LOW',    '6/7', '2/3', 'Sering lockout'),
    ('mti.chemicalanalyst', 0.3543, 'LOW',    '6/7', '2/3', 'Rasio login gagal'),
    ('muhammad.firmansyah', 0.3538, 'LOW',    '1/7', '2/3', 'Pelanggaran rule'),
]

# ── HELPERS ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_slide_header(doc, slide_num, title, subtitle=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ SECTION {slide_num} ]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(20)

    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(11)
            run.font.italic = True
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('-' * 70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
    doc.add_paragraph()

def add_bullet(doc, text, level=0, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(text[len(bold_part):])
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

def add_img(doc, fig, width=6.0):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    plt.close(fig)
    doc.add_paragraph()

def add_page_break(doc):
    doc.add_page_break()

# ── FIGURES ──────────────────────────────────────────────────────────

def fig_pipeline():
    fig, ax = plt.subplots(figsize=(12, 3))
    ax.axis('off')
    stages = [
        ("AD Log\nData", "#1A3C6E"),
        ("Neo4j\nKnowledge\nGraph", "#1565C0"),
        ("Rule-Based\nKnowledge\nEngine", "#1976D2"),
        ("Graph\nFeature\nExtraction", "#1E88E5"),
        ("Ensemble\nIF+LOF+EE", "#42A5F5"),
        ("SHAP\nExplainability", "#64B5F6"),
        ("Anomaly\nReport", "#90CAF9"),
    ]
    n = len(stages)
    w, h_box, gap = 1.4, 0.7, 0.15
    total = n * w + (n - 1) * gap
    x0 = (12 - total) / 2

    for i, (label, color) in enumerate(stages):
        x = x0 + i * (w + gap)
        rect = FancyBboxPatch((x, 0.15), w, h_box, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w / 2, 0.15 + h_box / 2, label, ha='center', va='center',
                fontsize=7.5, color='white', fontweight='bold', linespacing=1.4)
        if i < n - 1:
            ax.annotate('', xy=(x + w + gap, 0.15 + h_box / 2),
                        xytext=(x + w, 0.15 + h_box / 2),
                        arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        phase_label = ['1', '2', '3', '4', '5', '5.5', '6'][i]
        ax.text(x + w / 2, 0.08, f"Phase {phase_label}",
                ha='center', va='center', fontsize=6.5, color='#666')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 1.1)
    fig.tight_layout(pad=0)
    return fig

def fig_graph_schema():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.axis('off')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    nodes = {
        'User':       (5.0, 5.0, '#1A3C6E'),
        'Hostname':   (2.0, 3.5, '#1565C0'),
        'Server':     (8.0, 3.5, '#1976D2'),
        'IPAddress':  (2.0, 1.5, '#1E88E5'),
        'Group':      (8.0, 1.5, '#42A5F5'),
        'Event':      (5.0, 2.5, '#64B5F6'),
        'Service':    (5.0, 0.5, '#90CAF9'),
    }
    edges = [
        ('User', 'Hostname',  'LOGIN_FROM'),
        ('User', 'Server',    'AUTHENTICATED_VIA'),
        ('User', 'IPAddress', 'CONNECTED_FROM'),
        ('User', 'Group',     'MEMBER_OF'),
        ('Event','User',      'REFERENCES'),
        ('Hostname','IPAddress','USED_IP'),
        ('User', 'Service',   'USED_SERVICE'),
    ]
    for label, (x, y, color) in nodes.items():
        circle = plt.Circle((x, y), 0.45, color=color, zorder=3)
        ax.add_patch(circle)
        ax.text(x, y, label, ha='center', va='center', fontsize=7,
                color='white', fontweight='bold', zorder=4)
    for src, dst, rel in edges:
        x1, y1, _ = nodes[src]
        x2, y2, _ = nodes[dst]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#888', lw=1.5,
                                   connectionstyle='arc3,rad=0.1'), zorder=2)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my, rel, ha='center', va='center', fontsize=6,
                color='#333', style='italic',
                bbox=dict(boxstyle='round,pad=0.2', fc='#F0F4FF', ec='none'))
    ax.set_title('Neo4j Knowledge Graph - Node & Relationship Schema', fontsize=11,
                 fontweight='bold', color='#1A3C6E', pad=10)
    fig.tight_layout()
    return fig

def fig_severity_actual():
    """Distribusi severity AKTUAL dari hasil pipeline"""
    labels_order = ['HIGH', 'MEDIUM', 'LOW', 'NORMAL']
    sizes = [ANOMALY_DIST.get(k, 0) for k in labels_order]
    colors = ['#E53935', '#FB8C00', '#FDD835', '#43A047']
    explode = (0.1, 0.08, 0.02, 0)

    fig, ax = plt.subplots(figsize=(7, 4.5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=[f'{l}\n({s} users)' for l, s in zip(labels_order, sizes)],
        colors=colors, explode=explode, autopct='%1.1f%%',
        startangle=140, pctdistance=0.78, textprops={'fontsize': 9}
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title(f'Distribusi Severity Anomali (Aktual)\n{TOTAL_USERS} Users Dianalisis',
                 fontsize=10, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_ensemble_voting():
    """Distribusi voting ensemble"""
    labels = ['0 votes\n(normal)', '1 vote', '2 votes\n(anomali)', '3 votes\n(strong anomali)']
    sizes = [ENSEMBLE_VOTING['0_votes'], ENSEMBLE_VOTING['1_vote'],
             ENSEMBLE_VOTING['2_votes'], ENSEMBLE_VOTING['3_votes']]
    colors = ['#43A047', '#FDD835', '#FB8C00', '#E53935']

    fig, ax = plt.subplots(figsize=(8, 3.8))
    bars = ax.bar(labels, sizes, color=colors, edgecolor='white', linewidth=1)
    ax.set_ylabel('Jumlah User', fontsize=9)
    ax.set_title('Distribusi Ensemble Voting (3 Metode ML)',
                 fontsize=11, fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, sizes):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(sizes) * 0.02,
                f'{val}', ha='center', fontsize=9, fontweight='bold')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_yscale('log')
    fig.tight_layout()
    return fig

def fig_top_anomalies():
    """Bar chart top 10 anomalies aktual"""
    users = [a[0] for a in TOP_ANOMALIES]
    scores = [a[1] for a in TOP_ANOMALIES]
    severities = [a[2] for a in TOP_ANOMALIES]
    color_map = {'HIGH': '#E53935', 'MEDIUM': '#FB8C00', 'LOW': '#FDD835'}
    colors = [color_map[s] for s in severities]

    fig, ax = plt.subplots(figsize=(10, 4.5))
    bars = ax.barh(users[::-1], scores[::-1], color=colors[::-1],
                   edgecolor='white', linewidth=0.8)
    ax.set_xlabel('Anomaly Score', fontsize=9)
    ax.set_title('Top 10 Users dengan Anomaly Score Tertinggi',
                 fontsize=11, fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, scores[::-1]):
        ax.text(val + 0.005, bar.get_y() + bar.get_height() / 2,
                f'{val:.4f}', va='center', fontsize=8)
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    patches = [mpatches.Patch(color=c, label=s) for s, c in color_map.items()]
    ax.legend(handles=patches, fontsize=8, loc='lower right')
    fig.tight_layout()
    return fig

def fig_execution_done():
    """Status eksekusi: SEMUA SELESAI"""
    phases = [
        'Phase 1: Data Preparation',
        'Phase 2: Neo4j Ingestion',
        'Phase 3: Rule-Based Engine',
        'Phase 4: Feature Extraction',
        'Phase 5: Anomaly Detection',
        'Phase 5.5: SHAP',
        'Phase 6: Reporting',
    ]
    progress = [100] * 7

    fig, ax = plt.subplots(figsize=(9, 3.5))
    bars = ax.barh(phases, progress, color='#43A047', edgecolor='white',
                   linewidth=0.8, height=0.6)
    ax.set_xlim(0, 120)
    ax.set_xlabel('Progress (%)', fontsize=9)
    ax.set_title('Status Eksekusi Pipeline - SEMUA PHASE SELESAI',
                 fontsize=11, fontweight='bold', color='#1A3C6E')
    for bar in bars:
        ax.text(102, bar.get_y() + bar.get_height() / 2,
                'DONE', va='center', fontsize=9, fontweight='bold', color='#43A047')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_features():
    features = [
        'Host Diversity', 'Critical Server\nAccess Ratio', 'Failed Login\nRatio',
        'Shared Device\nRisk', 'IP Network\nRisk', 'Privilege\nLevel',
        'Graph\nConnectivity', 'Rule\nViolations'
    ]
    importance = [0.18, 0.14, 0.22, 0.10, 0.16, 0.08, 0.06, 0.06]
    colors = ['#1A3C6E', '#1565C0', '#1976D2', '#1E88E5',
              '#2196F3', '#42A5F5', '#64B5F6', '#90CAF9']
    fig, ax = plt.subplots(figsize=(10, 3.5))
    bars = ax.barh(features, importance, color=colors, edgecolor='white', linewidth=0.8)
    ax.set_xlabel('Relative Contribution to Anomaly Detection', fontsize=9)
    ax.set_title('8 Graph-Based Features - Indicative Importance', fontsize=11,
                 fontweight='bold', color='#1A3C6E')
    ax.set_xlim(0, 0.28)
    for bar, val in zip(bars, importance):
        ax.text(val + 0.003, bar.get_y() + bar.get_height() / 2,
                f'{val:.0%}', va='center', fontsize=8, color='#333')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_rules():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis('off')
    rules = [
        ('R001', 'Multi-host Login',           'Unique hosts > 3'),
        ('R002', 'Off-hours Login',            'Login di luar 08:00-18:00 > 10%'),
        ('R003', 'Shared Device',              'Device dipakai > 5 user'),
        ('R004', 'Critical Server Access',     'Akses ke DC / server CRITICAL'),
        ('R005', 'Failed Login Spike',         'Login gagal > 50'),
        ('R006', 'Unusual IP',                 'IP di luar Office/VPN'),
        ('R007', 'After-Hours Privileged',     'Admin akses DC di luar jam'),
    ]
    cols = ['Rule ID', 'Rule Name', 'Trigger Condition']
    col_w = [0.10, 0.40, 0.50]
    y_start = 0.92
    row_h = 0.115

    for j, (col, cw) in enumerate(zip(cols, col_w)):
        x = sum(col_w[:j])
        rect = FancyBboxPatch((x + 0.005, y_start), cw - 0.01, row_h,
                              boxstyle="round,pad=0.01", facecolor='#1A3C6E', edgecolor='none')
        ax.add_patch(rect)
        ax.text(x + cw / 2, y_start + row_h / 2, col, ha='center', va='center',
                color='white', fontsize=8.5, fontweight='bold', transform=ax.transAxes)
    for i, (rid, name, cond) in enumerate(rules):
        y = y_start - (i + 1) * row_h
        bg = '#EEF4FF' if i % 2 == 0 else '#F9FAFF'
        for j, (text, cw) in enumerate(zip([rid, name, cond], col_w)):
            x = sum(col_w[:j])
            rect = FancyBboxPatch((x + 0.005, y), cw - 0.01, row_h,
                                  boxstyle="round,pad=0.01", facecolor=bg, edgecolor='none')
            ax.add_patch(rect)
            clr = '#1A3C6E' if j == 0 else '#333'
            fw = 'bold' if j == 0 else 'normal'
            ax.text(x + cw / 2, y + row_h / 2, text, ha='center', va='center',
                    color=clr, fontsize=8, fontweight=fw, transform=ax.transAxes)
    ax.set_title('Rule-Based Knowledge Engine - 7 Domain Rules', fontsize=11,
                 fontweight='bold', color='#1A3C6E', pad=8)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    return fig

# ── BUILDER ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # SECTION 1: COVER
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN HASIL PIPELINE")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    title_p = doc.add_heading('Deteksi Anomali Active Directory\nBerbasis Knowledge Graph', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(24)

    doc.add_paragraph()
    sub = doc.add_paragraph('AD Log -> Neo4j Knowledge Graph -> Rule Engine ->\nGraph Features -> Ensemble (IF+LOF+EE) -> SHAP -> Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(11)
        run.font.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph(f'Mei 2026  |  Pipeline Selesai  |  {TOTAL_USERS} Users Dianalisis  |  45 Anomali Terdeteksi')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.5)
    add_page_break(doc)

    # SECTION 2: EXECUTIVE SUMMARY
    add_slide_header(doc, 2, "Executive Summary",
                     "Hasil deteksi anomali Active Directory")

    add_bullet(doc, "Hasil utama pipeline:", bold_part="Hasil utama pipeline:")
    add_bullet(doc, f"Total users dianalisis: {TOTAL_USERS}", level=1)
    add_bullet(doc, f"Anomali terdeteksi: 45 user (5.1% dari total)", level=1)
    add_bullet(doc, f"Severity HIGH: {ANOMALY_DIST.get('HIGH', 0)} user", level=1)
    add_bullet(doc, f"Severity MEDIUM: {ANOMALY_DIST.get('MEDIUM', 0)} user", level=1)
    add_bullet(doc, f"Severity LOW: {ANOMALY_DIST.get('LOW', 0)} user", level=1)
    add_bullet(doc, f"User normal: {ANOMALY_DIST.get('NORMAL', 0)} user", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Statistik anomaly score:", bold_part="Statistik anomaly score:")
    add_bullet(doc, f"Mean: {SCORE_STATS['mean']:.4f}", level=1)
    add_bullet(doc, f"Median: {SCORE_STATS['median']:.4f}", level=1)
    add_bullet(doc, f"Standard deviation: {SCORE_STATS['std']:.4f}", level=1)
    add_bullet(doc, f"Min: {SCORE_STATS['min']:.4f}  |  Max: {SCORE_STATS['max']:.4f}", level=1)

    doc.add_paragraph()
    add_bullet(doc, "User dengan score tertinggi:", bold_part="User dengan score tertinggi:")
    add_bullet(doc, f"mti.admin - Score 0.6800 (HIGH) - 7/7 rules dilanggar - 3/3 model setuju", level=1)
    add_bullet(doc, f"Top SHAP cause: Rasio login gagal", level=1)

    add_page_break(doc)

    # SECTION 3: ARCHITECTURE
    add_slide_header(doc, 3, "Arsitektur Pipeline",
                     "7-Phase Pipeline - AD Log ke Anomaly Output")

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Phase', 'Nama', 'Output Utama']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_bg(cell, '1A3C6E')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(10)

    rows_data = [
        ('Phase 1', 'Data Preparation', '1.8M events, 12 kolom, CSV bersih'),
        ('Phase 2', 'Neo4j Knowledge Graph', f'{TOTAL_USERS} User nodes, 7 node types, 8 relationships'),
        ('Phase 3', 'Rule-Based Knowledge Engine', '7 domain rules dievaluasi via Cypher'),
        ('Phase 4', 'Graph Feature Extraction', '8 fitur per user, output CSV'),
        ('Phase 5', 'Ensemble Anomaly Detection', 'IF + LOF + EE -> score + severity'),
        ('Phase 5.5', 'SHAP Explainability', 'SHAP values, top causes per user'),
        ('Phase 6', 'Reporting', 'Text report, JSON detail, statistics, DOCX'),
    ]
    for i, (phase, name, output) in enumerate(rows_data):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        for j, text in enumerate([phase, name, output]):
            cell = table.cell(i + 1, j)
            cell.text = text
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if j == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.2)
    add_page_break(doc)

    # SECTION 4: DATA OVERVIEW
    add_slide_header(doc, 4, "Phase 1-2 - Data & Knowledge Graph",
                     "Sumber data dan struktur graph yang dibangun")

    add_bullet(doc, "Statistik dataset input:", bold_part="Statistik dataset input:")
    add_bullet(doc, "Total events: 1.833.352 baris dari unified_logon_events.csv", level=1)
    add_bullet(doc, f"Total users unik: {TOTAL_USERS} (setelah ingestion)", level=1)
    add_bullet(doc, "Total hostnames: 1.273 device unik", level=1)
    add_bullet(doc, "Total IP addresses: 1.275 IP unik", level=1)
    add_bullet(doc, "Servers: 12 (termasuk Domain Controllers)", level=1)
    add_bullet(doc, "Groups: 3 (User, Admin, Service)", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Node & Relationship di Neo4j:", bold_part="Node & Relationship di Neo4j:")
    add_img(doc, fig_graph_schema(), width=6.0)
    add_page_break(doc)

    # SECTION 5: RULES
    add_slide_header(doc, 5, "Phase 3 - Rule-Based Engine",
                     "7 domain rules dievaluasi langsung di Neo4j via Cypher")

    add_img(doc, fig_rules(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Output per user setelah Phase 3:")
    add_bullet(doc, "u.rule_R001_violation = true/false", level=1)
    add_bullet(doc, "u.rule_violations = total rules dilanggar (0-7)", level=1)
    add_bullet(doc, "u.max_rule_severity = HIGH / MEDIUM / LOW", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Distribusi pelanggaran rule (aktual):",
               bold_part="Distribusi pelanggaran rule (aktual):")
    add_bullet(doc, f"User dengan 0 violation: {STATS['rule_violations_stats']['users_with_0']}", level=1)
    add_bullet(doc, f"User dengan 1-2 violations: {STATS['rule_violations_stats']['users_with_1_2']}", level=1)
    add_bullet(doc, f"User dengan 3+ violations: {STATS['rule_violations_stats']['users_with_3_plus']}", level=1)
    add_bullet(doc, f"Rata-rata: {STATS['rule_violations_stats']['mean']:.2f} violations per user", level=1)

    add_page_break(doc)

    # SECTION 6: FEATURES
    add_slide_header(doc, 6, "Phase 4 - Graph Feature Extraction",
                     "8 fitur diekstrak dari relasi graph - bukan dari kolom CSV")

    add_img(doc, fig_features(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Output: data/phase4_graph_features.csv")
    add_bullet(doc, f"{TOTAL_USERS} rows x 8 kolom, siap input ke Phase 5", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Kenapa fitur dari graph, bukan dari CSV langsung?")
    add_bullet(doc, "CSV hanya event individual - tidak ada konteks relasi", level=1)
    add_bullet(doc, "Graph menghitung HUBUNGAN: berapa device, server, IP unik per user", level=1)
    add_bullet(doc, "Graph Connectivity (degree centrality) tidak bisa dihitung dari CSV", level=1)
    add_bullet(doc, "Rule violations terintegrasi sebagai feature (bukan post-processing)", level=1)

    add_page_break(doc)

    # SECTION 7: ANOMALY DETECTION RESULTS
    add_slide_header(doc, 7, "Phase 5 - Hasil Deteksi Anomali",
                     "Distribusi severity dan voting ensemble (data aktual)")

    add_img(doc, fig_severity_actual(), width=5.5)
    doc.add_paragraph()
    add_img(doc, fig_ensemble_voting(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Insight dari distribusi:")
    add_bullet(doc, f"{ENSEMBLE_VOTING['3_votes']} user dengan 3/3 votes (strong anomali)", level=1)
    add_bullet(doc, f"{ENSEMBLE_VOTING['2_votes']} user dengan 2/3 votes (anomali)", level=1)
    add_bullet(doc, f"Total {ENSEMBLE_VOTING['3_votes'] + ENSEMBLE_VOTING['2_votes']} user terdeteksi anomali oleh ensemble", level=1)
    add_bullet(doc, f"Mayoritas ({ENSEMBLE_VOTING['0_votes']} user) tidak ter-flag oleh model manapun (normal)", level=1)

    add_page_break(doc)

    # SECTION 8: TOP ANOMALIES
    add_slide_header(doc, 8, "Top 10 Anomalies Detected",
                     "User dengan anomaly score tertinggi - hasil aktual")

    add_img(doc, fig_top_anomalies(), width=6.5)

    doc.add_paragraph()
    table_top = doc.add_table(rows=len(TOP_ANOMALIES) + 1, cols=6)
    table_top.style = 'Table Grid'
    headers = ['Rank', 'Username', 'Score', 'Severity', 'Rules', 'Top SHAP Cause']
    for j, h in enumerate(headers):
        c = table_top.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)

    sev_color = {'HIGH': 'E53935', 'MEDIUM': 'FB8C00', 'LOW': 'FDD835'}
    for i, (uname, score, sev, rules, votes, cause) in enumerate(TOP_ANOMALIES):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        cells = [str(i + 1), uname, f'{score:.4f}', sev, rules, cause]
        for j, text in enumerate(cells):
            c = table_top.cell(i + 1, j)
            c.text = text
            if j == 3:
                set_cell_bg(c, sev_color[sev])
            else:
                set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
                if j == 3:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                if j == 0:
                    run.font.bold = True

    add_page_break(doc)

    # SECTION 9: SHAP CAUSES
    add_slide_header(doc, 9, "Phase 5.5 - SHAP Explainability",
                     "Mengapa setiap user dianggap anomali")

    add_bullet(doc, "Top SHAP causes pada anomali (dari hasil aktual):",
               bold_part="Top SHAP causes pada anomali (dari hasil aktual):")
    causes = {}
    for _, _, _, _, _, cause in TOP_ANOMALIES:
        causes[cause] = causes.get(cause, 0) + 1
    for cause, count in sorted(causes.items(), key=lambda x: -x[1]):
        add_bullet(doc, f'"{cause}": muncul pada {count} dari top 10 anomalies', level=1)

    doc.add_paragraph()
    add_bullet(doc, "Metode yang digunakan:")
    add_bullet(doc, "shap.TreeExplainer pada model Isolation Forest", level=1)
    add_bullet(doc, "Menghitung kontribusi tiap fitur ke anomaly score", level=1)
    add_bullet(doc, f"Output: SHAP value per fitur per user ({TOTAL_USERS} users x 8 features)", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Interpretasi SHAP:")
    add_bullet(doc, "Nilai positif = fitur ini menaikkan anomaly score", level=1)
    add_bullet(doc, "Nilai negatif = fitur ini menurunkan anomaly score", level=1)
    add_bullet(doc, "Nilai absolut terbesar = faktor paling dominan", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Manfaat untuk audit:")
    add_bullet(doc, "Auditor bisa jelaskan: 'mti.admin anomali karena rasio login gagal tinggi'", level=1)
    add_bullet(doc, "Bukan sekedar label anomali tanpa bukti kuantitatif", level=1)
    add_bullet(doc, "Setiap temuan punya audit trail yang traceable", level=1)

    add_page_break(doc)

    # SECTION 10: KEY FINDINGS
    add_slide_header(doc, 10, "Key Findings & Recommendations",
                     "Temuan utama dan rekomendasi tindak lanjut")

    h_find = doc.add_heading("Key Findings", level=2)
    for run in h_find.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_bullet(doc, "Brute Force Risk: 712 user dengan failure_ratio > 20%")
    add_bullet(doc, "Shared Device Risk: 686 user mengakses shared device")
    add_bullet(doc, "Top user (mti.admin) melanggar 7/7 rules dengan score 0.68")
    add_bullet(doc, f"Hanya {ENSEMBLE_VOTING['3_votes']} user terdeteksi oleh 3/3 model (strong anomali)")

    doc.add_paragraph()
    h_rec = doc.add_heading("Rekomendasi Tindak Lanjut", level=2)
    for run in h_rec.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_bullet(doc, "Immediate (HIGH/MEDIUM):", bold_part="Immediate (HIGH/MEDIUM):")
    add_bullet(doc, "Reset password mti.admin dan mti.sysadmin", level=1)
    add_bullet(doc, "Audit privileged access logs untuk top 10 user", level=1)
    add_bullet(doc, "Investigasi pola login gagal tinggi", level=1)

    add_bullet(doc, "Short-term (LOW severity):", bold_part="Short-term (LOW severity):")
    add_bullet(doc, "Enable MFA untuk admin account", level=1)
    add_bullet(doc, "Monitor shared device usage", level=1)
    add_bullet(doc, "Review group memberships", level=1)

    add_bullet(doc, "Long-term:", bold_part="Long-term:")
    add_bullet(doc, "Implement risk-based adaptive authentication", level=1)
    add_bullet(doc, "Strengthen network segmentation", level=1)
    add_bullet(doc, "Deploy continuous behavioral analytics", level=1)

    add_page_break(doc)

    # SECTION 11: EXECUTION STATUS
    add_slide_header(doc, 11, "Status Eksekusi",
                     "Semua phase telah berhasil dijalankan")

    add_img(doc, fig_execution_done(), width=6.0)

    doc.add_paragraph()
    status_data = [
        ('Phase 1', 'Data Preparation', 'DONE', '1.8M events validated, CSV unified'),
        ('Phase 2', 'Neo4j Ingestion', 'DONE', f'{TOTAL_USERS} users, 1.273 hostnames, 1.275 IPs'),
        ('Phase 3', 'Rule-Based Engine', 'DONE', '7 rules diterapkan, mean 3.58 violations/user'),
        ('Phase 4', 'Feature Extraction', 'DONE', f'{TOTAL_USERS} users x 8 features extracted'),
        ('Phase 5', 'Anomaly Detection', 'DONE', '45 anomali terdeteksi (5.1%)'),
        ('Phase 5.5', 'SHAP Explainability', 'DONE', 'Top causes per user di-export'),
        ('Phase 6', 'Reporting', 'DONE', 'TXT, JSON, DOCX, statistics'),
    ]
    t5 = doc.add_table(rows=len(status_data) + 1, cols=4)
    t5.style = 'Table Grid'
    for j, h in enumerate(['Phase', 'Nama', 'Status', 'Hasil']):
        c = t5.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)
    for i, (ph, name, status, note) in enumerate(status_data):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        for j, text in enumerate([ph, name, status, note]):
            c = t5.cell(i + 1, j)
            c.text = text
            set_cell_bg(c, bg if j != 2 else '43A047')
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                if j == 2:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True

    add_page_break(doc)

    # SECTION 12: OUTPUT FILES
    add_slide_header(doc, 12, "File Output",
                     "Daftar file yang dihasilkan pipeline")

    outputs = [
        ('output/anomaly_detection_report.txt', 'Laporan teks untuk auditor'),
        ('output/anomaly_detection_detailed.json', 'Detail per user (SHAP + rules)'),
        ('output/anomaly_statistics.json', 'Statistik distribusi anomali'),
        ('output/AD_Anomaly_Detection_Report.docx', 'Laporan DOCX (file ini)'),
        ('data/phase4_graph_features.csv', f'{TOTAL_USERS} user x 8 features'),
        ('data/phase5_anomaly_results.csv', 'Score & severity semua user'),
        ('data/phase55_shap_values.csv', f'SHAP values: {TOTAL_USERS} user x 8 features'),
        ('data/phase55_shap_anomalies.csv', '45 anomali dengan top causes'),
        ('models/isolation_forest_model.pkl', 'Trained IF model'),
        ('models/lof_model.pkl', 'Trained LOF model'),
        ('models/elliptic_envelope_model.pkl', 'Trained EE model'),
        ('models/feature_scaler.pkl', 'StandardScaler'),
    ]
    table4 = doc.add_table(rows=len(outputs) + 1, cols=2)
    table4.style = 'Table Grid'
    for j, h in enumerate(['File Path', 'Isi']):
        c = table4.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)
    for i, (fname, desc) in enumerate(outputs):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        c1 = table4.cell(i + 1, 0)
        c2 = table4.cell(i + 1, 1)
        c1.text = fname
        c2.text = desc
        for c in [c1, c2]:
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
        for run in c1.paragraphs[0].runs:
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    add_divider(doc)

    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final.add_run("AD Log -> Neo4j -> Rules -> Features -> Ensemble -> SHAP -> Report")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("Pipeline Selesai - Mei 2026")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rf.font.italic = True

    out_path = 'output/AD_Anomaly_Detection_Report_v2.docx'
    doc.save(out_path)
    print(f"[OK] DOCX saved: {out_path}")
    return out_path


if __name__ == '__main__':
    print("Generating updated report...")
    path = build_document()
    print(f"Done: {path}")
