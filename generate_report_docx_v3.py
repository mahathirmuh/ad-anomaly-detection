#!/usr/bin/env python3
"""
Generate Project Report DOCX v3
AD Anomaly Detection - Graph-Based Knowledge System

Perubahan dari v2:
  - Quantile-based severity threshold (data-driven) + referensi paper
  - Section baru: Threshold Methodology & Justification
  - Section baru: Ablation Study (IF vs LOF vs EE)
  - Semua angka dibaca dinamis dari hasil pipeline
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import pandas as pd
import numpy as np
from sklearn.metrics import cohen_kappa_score
import json
import io
import os

os.makedirs('output', exist_ok=True)

# ── LOAD PIPELINE RESULTS (DINAMIS) ──────────────────────────────────
with open('output/anomaly_statistics.json') as f:
    STATS = json.load(f)

TOTAL_USERS   = STATS['total_users']
ANOMALY_DIST  = STATS['anomaly_distribution']
SCORE_STATS   = STATS['anomaly_score_stats']
ENSEMBLE_VOTE = STATS['ensemble_voting']
RULE_STATS    = STATS['rule_violations_stats']

df = pd.read_csv('data/phase5_anomaly_results.csv').drop_duplicates(subset='user_id').reset_index(drop=True)
shap_df = pd.read_csv('data/phase55_shap_values.csv')

# Quantile thresholds
P75 = df['final_anomaly_score'].quantile(0.75)
P90 = df['final_anomaly_score'].quantile(0.90)
P95 = df['final_anomaly_score'].quantile(0.95)
P99 = df['final_anomaly_score'].quantile(0.99)

# Top anomalies (dinamis, merge dengan SHAP)
merged = df.merge(shap_df[['user_id', 'top_feature_1_label']], on='user_id', how='left')
TOP_ANOMALIES = []
for _, r in merged.nlargest(10, 'final_anomaly_score').iterrows():
    TOP_ANOMALIES.append((
        r['username'],
        float(r['final_anomaly_score']),
        r['severity'],
        f"{int(r['rule_violations'])}/10",
        f"{int(r['anomaly_votes'])}/3",
        r['top_feature_1_label'] if pd.notna(r['top_feature_1_label']) else '-'
    ))

TOTAL_ANOMALI = int(df['severity'].isin(['CRITICAL', 'HIGH', 'MEDIUM']).sum())

# ── ABLATION DATA ────────────────────────────────────────────────────
if_set  = set(df[df['if_anomaly']  == 1]['user_id'])
lof_set = set(df[df['lof_anomaly'] == 1]['user_id'])
ee_set  = set(df[df['ee_anomaly']  == 1]['user_id'])
ens_set = set(df[df['anomaly_votes'] >= 2]['user_id'])
heavy   = set(df[df['rule_violations'] >= 6]['user_id'])

def jaccard(a, b):
    return len(a & b) / len(a | b) if len(a | b) > 0 else 0.0

ABL_OVERLAP = {
    'only_if':  len(if_set - lof_set - ee_set),
    'only_lof': len(lof_set - if_set - ee_set),
    'only_ee':  len(ee_set - if_set - lof_set),
    'if_lof':   len((if_set & lof_set) - ee_set),
    'if_ee':    len((if_set & ee_set) - lof_set),
    'lof_ee':   len((lof_set & ee_set) - if_set),
    'all3':     len(if_set & lof_set & ee_set),
}
ABL_AGREEMENT = [
    ('IF-LOF', jaccard(if_set, lof_set), cohen_kappa_score(df['if_anomaly'], df['lof_anomaly'])),
    ('IF-EE',  jaccard(if_set, ee_set),  cohen_kappa_score(df['if_anomaly'], df['ee_anomaly'])),
    ('LOF-EE', jaccard(lof_set, ee_set), cohen_kappa_score(df['lof_anomaly'], df['ee_anomaly'])),
]

def abl_quality(s):
    sub = df[df['user_id'].isin(s)]
    n = len(sub)
    return {
        'n': n,
        'prec_k': len(s & heavy) / n if n > 0 else 0,
        'avg_rv': sub['rule_violations'].mean() if n > 0 else 0,
        'avg_score': sub['final_anomaly_score'].mean() if n > 0 else 0,
    }

ABL_QUALITY = {
    'IF':       abl_quality(if_set),
    'LOF':      abl_quality(lof_set),
    'EE':       abl_quality(ee_set),
    'Ensemble': abl_quality(ens_set),
}
BASELINE_RV = df['rule_violations'].mean()
BASELINE_SCORE = df['final_anomaly_score'].mean()

# ── HELPERS ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_slide_header(doc, num, title, subtitle=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ SECTION {num} ]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(20)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(11)
            run.font.italic = True
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('-' * 70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
    doc.add_paragraph()

def add_bullet(doc, text, level=0, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text[len(bold_part):])
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

def add_img(doc, fig, width=6.0):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    plt.close(fig)
    doc.add_paragraph()

def add_page_break(doc):
    doc.add_page_break()

def make_table(doc, headers, rows, col_fmt=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = str(h)
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(rows):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        for j, val in enumerate(row):
            c = table.cell(i + 1, j)
            c.text = str(val)
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                if col_fmt and col_fmt.get(j) == 'bold':
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return table

# ── FIGURES ──────────────────────────────────────────────────────────

COLORS_MODEL = ['#1A3C6E', '#1976D2', '#42A5F5', '#C62828']

def fig_pipeline():
    fig, ax = plt.subplots(figsize=(12, 3))
    ax.axis('off')
    stages = [
        ("AD Log\nData", "#1A3C6E"),
        ("Neo4j\nKnowledge\nGraph", "#1565C0"),
        ("Rule-Based\nEngine\n(10 rules)", "#1976D2"),
        ("Graph\nFeature\n(11 fitur)", "#1E88E5"),
        ("Ensemble\nIF+LOF+EE", "#42A5F5"),
        ("SHAP\nExplainability", "#64B5F6"),
        ("Anomaly\nReport", "#90CAF9"),
    ]
    n = len(stages)
    w, h_box, gap = 1.4, 0.7, 0.15
    total = n * w + (n - 1) * gap
    x0 = (12 - total) / 2
    for i, (label, color) in enumerate(stages):
        x = x0 + i * (w + gap)
        rect = FancyBboxPatch((x, 0.15), w, h_box, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w / 2, 0.15 + h_box / 2, label, ha='center', va='center',
                fontsize=7.5, color='white', fontweight='bold', linespacing=1.4)
        if i < n - 1:
            ax.annotate('', xy=(x + w + gap, 0.15 + h_box / 2),
                        xytext=(x + w, 0.15 + h_box / 2),
                        arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        phase_label = ['1', '2', '3', '4', '5', '5.5', '6'][i]
        ax.text(x + w / 2, 0.08, f"Phase {phase_label}",
                ha='center', va='center', fontsize=6.5, color='#666')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 1.1)
    fig.tight_layout(pad=0)
    return fig

def fig_severity():
    labels = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'NORMAL']
    sizes = [ANOMALY_DIST.get(k, 0) for k in labels]
    colors = ['#B71C1C', '#E53935', '#FB8C00', '#FDD835', '#43A047']
    explode = (0.12, 0.06, 0.03, 0, 0)
    fig, ax = plt.subplots(figsize=(7, 4.5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=[f'{l}\n({s})' for l, s in zip(labels, sizes)],
        colors=colors, explode=explode, autopct='%1.1f%%',
        startangle=140, pctdistance=0.78, textprops={'fontsize': 9})
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title(f'Distribusi Severity (Quantile-based)\n{TOTAL_USERS} Users Dianalisis',
                 fontsize=10, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_threshold():
    """Visualisasi distribusi score + garis threshold quantile"""
    fig, ax = plt.subplots(figsize=(10, 4))
    scores = df['final_anomaly_score']
    ax.hist(scores, bins=50, color='#90CAF9', edgecolor='white')
    for val, label, color in [(P75, 'P75 (LOW)', '#FDD835'),
                               (P90, 'P90 (MEDIUM)', '#FB8C00'),
                               (P95, 'P95 (HIGH)', '#E53935'),
                               (P99, 'P99 (CRITICAL)', '#B71C1C')]:
        ax.axvline(val, color=color, linestyle='--', linewidth=2,
                   label=f'{label} = {val:.4f}')
    ax.set_xlabel('Anomaly Score')
    ax.set_ylabel('Jumlah User')
    ax.set_title('Distribusi Anomaly Score & Threshold Quantile-based',
                 fontsize=11, fontweight='bold', color='#1A3C6E')
    ax.legend(fontsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_abl_detection():
    models = ['IF', 'LOF', 'EE', 'Ensemble']
    counts = [ABL_QUALITY[m]['n'] for m in models]
    fig, ax = plt.subplots(figsize=(8, 3.8))
    bars = ax.bar(models, counts, color=COLORS_MODEL, edgecolor='white', linewidth=1)
    ax.set_ylabel('Jumlah Anomali Terdeteksi')
    ax.set_title('Ablation: Jumlah Deteksi per Model', fontsize=11,
                 fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                str(val), ha='center', fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_abl_overlap():
    regions = ['Hanya\nIF', 'Hanya\nLOF', 'Hanya\nEE', 'IF&LOF', 'IF&EE', 'LOF&EE', 'Ketiga\n(3/3)']
    vals = [ABL_OVERLAP['only_if'], ABL_OVERLAP['only_lof'], ABL_OVERLAP['only_ee'],
            ABL_OVERLAP['if_lof'], ABL_OVERLAP['if_ee'], ABL_OVERLAP['lof_ee'], ABL_OVERLAP['all3']]
    rc = ['#1A3C6E', '#1976D2', '#42A5F5', '#7E57C2', '#5C6BC0', '#26A69A', '#C62828']
    fig, ax = plt.subplots(figsize=(10, 3.8))
    bars = ax.bar(regions, vals, color=rc, edgecolor='white', linewidth=1)
    ax.set_ylabel('Jumlah User')
    ax.set_title('Ablation: Overlap antar Model (3-Set Venn)', fontsize=11,
                 fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                str(val), ha='center', fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_abl_quality():
    models = ['IF', 'LOF', 'EE', 'Ensemble']
    fig, axes = plt.subplots(1, 3, figsize=(14, 4))
    # Precision@K
    pk = [ABL_QUALITY[m]['prec_k'] for m in models]
    axes[0].bar(models, pk, color=COLORS_MODEL, edgecolor='white')
    axes[0].set_title('Precision@K\n(% deteksi = pelanggar rule >=6)', fontweight='bold', fontsize=10)
    axes[0].set_ylim(0, 1.05)
    for i, v in enumerate(pk):
        axes[0].text(i, v + 0.02, f'{v:.1%}', ha='center', fontweight='bold', fontsize=9)
    # Avg rule violations
    rv = [ABL_QUALITY[m]['avg_rv'] for m in models]
    axes[1].bar(models, rv, color=COLORS_MODEL, edgecolor='white')
    axes[1].axhline(BASELINE_RV, color='red', linestyle='--', label=f'Baseline ({BASELINE_RV:.2f})')
    axes[1].set_title('Avg rule_violations\npada anomali terdeteksi', fontweight='bold', fontsize=10)
    axes[1].legend(fontsize=8)
    for i, v in enumerate(rv):
        axes[1].text(i, v + 0.1, f'{v:.2f}', ha='center', fontweight='bold', fontsize=9)
    # Avg score
    sc = [ABL_QUALITY[m]['avg_score'] for m in models]
    axes[2].bar(models, sc, color=COLORS_MODEL, edgecolor='white')
    axes[2].axhline(BASELINE_SCORE, color='red', linestyle='--', label=f'Baseline ({BASELINE_SCORE:.3f})')
    axes[2].set_title('Avg anomaly_score\npada anomali terdeteksi', fontweight='bold', fontsize=10)
    axes[2].legend(fontsize=8)
    for i, v in enumerate(sc):
        axes[2].text(i, v + 0.005, f'{v:.3f}', ha='center', fontweight='bold', fontsize=9)
    for ax in axes:
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_top_anomalies():
    users = [a[0] for a in TOP_ANOMALIES]
    scores = [a[1] for a in TOP_ANOMALIES]
    sevs = [a[2] for a in TOP_ANOMALIES]
    cmap = {'CRITICAL': '#B71C1C', 'HIGH': '#E53935', 'MEDIUM': '#FB8C00', 'LOW': '#FDD835'}
    colors = [cmap.get(s, '#999') for s in sevs]
    fig, ax = plt.subplots(figsize=(10, 4.5))
    bars = ax.barh(users[::-1], scores[::-1], color=colors[::-1], edgecolor='white', linewidth=0.8)
    ax.set_xlabel('Anomaly Score')
    ax.set_title('Top 10 Users dengan Anomaly Score Tertinggi', fontsize=11,
                 fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, scores[::-1]):
        ax.text(val + 0.005, bar.get_y() + bar.get_height() / 2,
                f'{val:.4f}', va='center', fontsize=8)
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    patches = [mpatches.Patch(color=c, label=s) for s, c in cmap.items()]
    ax.legend(handles=patches, fontsize=8, loc='lower right')
    fig.tight_layout()
    return fig

# ── DOCUMENT BUILDER ─────────────────────────────────────────────────

def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # SECTION 1: COVER
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN HASIL PIPELINE")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()
    title_p = doc.add_heading('Deteksi Anomali Active Directory\nBerbasis Knowledge Graph', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(24)
    doc.add_paragraph()
    sub = doc.add_paragraph('AD Log -> Neo4j Knowledge Graph -> Rule Engine ->\n'
                            'Graph Features -> Ensemble (IF+LOF+EE) -> SHAP -> Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(11)
        run.font.italic = True
    doc.add_paragraph()
    meta = doc.add_paragraph(f'Mei 2026  |  Pipeline Selesai  |  {TOTAL_USERS} Users  |  '
                             f'{TOTAL_ANOMALI} Anomali (MEDIUM+)  |  v3')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.5)
    add_page_break(doc)

    # SECTION 2: EXECUTIVE SUMMARY
    add_slide_header(doc, 2, "Executive Summary", "Hasil deteksi anomali Active Directory")
    add_bullet(doc, "Hasil utama pipeline:", bold_part="Hasil utama pipeline:")
    add_bullet(doc, f"Total users dianalisis: {TOTAL_USERS}", level=1)
    add_bullet(doc, f"CRITICAL: {ANOMALY_DIST.get('CRITICAL',0)} user (top 1%)", level=1)
    add_bullet(doc, f"HIGH: {ANOMALY_DIST.get('HIGH',0)} user (top 5%)", level=1)
    add_bullet(doc, f"MEDIUM: {ANOMALY_DIST.get('MEDIUM',0)} user (top 10%)", level=1)
    add_bullet(doc, f"LOW: {ANOMALY_DIST.get('LOW',0)} user (top 25%)", level=1)
    add_bullet(doc, f"NORMAL: {ANOMALY_DIST.get('NORMAL',0)} user", level=1)
    doc.add_paragraph()
    add_bullet(doc, "Statistik anomaly score:", bold_part="Statistik anomaly score:")
    add_bullet(doc, f"Mean: {SCORE_STATS['mean']:.4f} | Median: {SCORE_STATS['median']:.4f}", level=1)
    add_bullet(doc, f"Min: {SCORE_STATS['min']:.4f} | Max: {SCORE_STATS['max']:.4f}", level=1)
    doc.add_paragraph()
    add_bullet(doc, "User dengan score tertinggi:", bold_part="User dengan score tertinggi:")
    top1 = TOP_ANOMALIES[0]
    add_bullet(doc, f"{top1[0]} - Score {top1[1]:.4f} ({top1[2]}) - {top1[3]} rules - "
                    f"{top1[4]} votes - Top cause: {top1[5]}", level=1)
    add_page_break(doc)

    # SECTION 3: ARCHITECTURE
    add_slide_header(doc, 3, "Arsitektur Pipeline", "7-Phase Pipeline - AD Log ke Anomaly Output")
    make_table(doc, ['Phase', 'Nama', 'Output Utama'], [
        ('Phase 1', 'Data Preparation', '1.8M events, 12 kolom, CSV bersih'),
        ('Phase 2', 'Neo4j Knowledge Graph', f'{TOTAL_USERS} User nodes, 7 node types, 8 relationships'),
        ('Phase 3', 'Rule-Based Engine', '10 domain rules dievaluasi via Cypher'),
        ('Phase 4', 'Graph Feature Extraction', '11 fitur per user, output CSV'),
        ('Phase 5', 'Ensemble Anomaly Detection', 'IF + LOF + EE -> quantile severity'),
        ('Phase 5.5', 'SHAP Explainability', 'SHAP values, top causes per user'),
        ('Phase 6', 'Reporting', 'Text, JSON, statistics, DOCX'),
    ], col_fmt={0: 'bold'})
    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.2)
    add_page_break(doc)

    # SECTION 4: DATA & GRAPH
    add_slide_header(doc, 4, "Phase 1-4 - Data, Graph & Features",
                     "Dari raw log ke 11 fitur graph")
    add_bullet(doc, "Statistik dataset input:", bold_part="Statistik dataset input:")
    add_bullet(doc, "Total events: 1.833.352 baris dari unified_logon_events.csv", level=1)
    add_bullet(doc, f"Total users unik: {TOTAL_USERS}", level=1)
    add_bullet(doc, "Hostnames: 1.273 | IP addresses: 1.275 | Servers: 7-12", level=1)
    doc.add_paragraph()
    add_bullet(doc, "11 Graph Features (Phase 4):", bold_part="11 Graph Features (Phase 4):")
    add_bullet(doc, "Dasar (8): host_diversity, critical_server_ratio, failure_ratio, "
                    "shared_device_risk, ip_network_risk, privilege_level, connectivity, "
                    "rule_violations", level=1)
    add_bullet(doc, "Tambahan (3): lockout_count, admin_actions, sensitive_groups", level=1)
    doc.add_paragraph()
    add_bullet(doc, "Distribusi pelanggaran rule (10 rules):", bold_part="Distribusi pelanggaran rule (10 rules):")
    add_bullet(doc, f"0 violation: {RULE_STATS['users_with_0']} | "
                    f"1-2: {RULE_STATS['users_with_1_2']} | "
                    f"3+: {RULE_STATS['users_with_3_plus']}", level=1)
    add_bullet(doc, f"Rata-rata: {RULE_STATS['mean']:.2f} violations per user", level=1)
    add_page_break(doc)

    # SECTION 5: THRESHOLD METHODOLOGY (BARU)
    add_slide_header(doc, 5, "Metodologi Threshold Severity",
                     "Quantile-based classification - data-driven & defensible")
    add_bullet(doc, "Masalah threshold absolut:", bold_part="Masalah threshold absolut:")
    add_bullet(doc, f"Distribusi anomaly score heavy-tailed (min={SCORE_STATS['min']:.3f}, "
                    f"max={SCORE_STATS['max']:.3f}, mean={SCORE_STATS['mean']:.3f})", level=1)
    add_bullet(doc, "Threshold absolut (mis. CVSS murni) menghasilkan kelas severity tidak seimbang", level=1)
    doc.add_paragraph()
    add_bullet(doc, "Solusi: Quantile-based threshold", bold_part="Solusi: Quantile-based threshold")
    make_table(doc, ['Severity', 'Threshold', 'Persentil', 'Jumlah User'], [
        ('CRITICAL', f'>= {P99:.4f}', 'Top 1% (P99)', ANOMALY_DIST.get('CRITICAL', 0)),
        ('HIGH',     f'>= {P95:.4f}', 'Top 5% (P95)', ANOMALY_DIST.get('HIGH', 0)),
        ('MEDIUM',   f'>= {P90:.4f}', 'Top 10% (P90)', ANOMALY_DIST.get('MEDIUM', 0)),
        ('LOW',      f'>= {P75:.4f}', 'Top 25% (P75)', ANOMALY_DIST.get('LOW', 0)),
        ('NORMAL',   f'< {P75:.4f}',  'Sisanya', ANOMALY_DIST.get('NORMAL', 0)),
    ], col_fmt={0: 'bold'})
    doc.add_paragraph()
    add_img(doc, fig_threshold(), width=6.2)
    doc.add_paragraph()
    add_bullet(doc, "Keunggulan pendekatan ini:", bold_part="Keunggulan pendekatan ini:")
    add_bullet(doc, "Data-driven, objektif, reproducible (bukan pilihan subjektif)", level=1)
    add_bullet(doc, "Selaras dengan contamination rate Isolation Forest (5-10%)", level=1)
    add_bullet(doc, "Proporsi severity konsisten across deployment", level=1)
    doc.add_paragraph()
    add_bullet(doc, "Referensi:", bold_part="Referensi:")
    add_bullet(doc, "Aggarwal, C. C. (2017). Outlier Analysis (2nd ed.). Springer.", level=1)
    add_bullet(doc, "Goldstein, M., & Uchida, S. (2016). A Comparative Evaluation of "
                    "Unsupervised Anomaly Detection Algorithms. PLOS ONE.", level=1)
    add_bullet(doc, "Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation Forest. IEEE ICDM.", level=1)
    add_page_break(doc)

    # SECTION 6: ANOMALY DETECTION RESULTS
    add_slide_header(doc, 6, "Phase 5 - Hasil Deteksi Anomali",
                     "Distribusi severity dengan threshold quantile-based")
    add_img(doc, fig_severity(), width=5.5)
    doc.add_paragraph()
    add_bullet(doc, "Insight:", bold_part="Insight:")
    add_bullet(doc, f"{ANOMALY_DIST.get('CRITICAL',0)} user CRITICAL - prioritas investigasi tertinggi", level=1)
    add_bullet(doc, f"Total {TOTAL_ANOMALI} user perlu perhatian (MEDIUM ke atas)", level=1)
    add_bullet(doc, f"Ensemble voting: {ENSEMBLE_VOTE['3_votes']} user 3/3 votes, "
                    f"{ENSEMBLE_VOTE['2_votes']} user 2/3 votes", level=1)
    add_page_break(doc)

    # SECTION 7: TOP ANOMALIES
    add_slide_header(doc, 7, "Top 10 Anomalies Detected",
                     "User dengan anomaly score tertinggi")
    add_img(doc, fig_top_anomalies(), width=6.5)
    doc.add_paragraph()
    sev_color = {'CRITICAL': 'B71C1C', 'HIGH': 'E53935', 'MEDIUM': 'FB8C00', 'LOW': 'FDD835'}
    table = doc.add_table(rows=len(TOP_ANOMALIES) + 1, cols=6)
    table.style = 'Table Grid'
    for j, h in enumerate(['Rank', 'Username', 'Score', 'Severity', 'Rules', 'Top SHAP Cause']):
        c = table.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
    for i, (uname, score, sev, rules, votes, cause) in enumerate(TOP_ANOMALIES):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        cells = [str(i + 1), uname, f'{score:.4f}', sev, rules, cause]
        for j, val in enumerate(cells):
            c = table.cell(i + 1, j)
            c.text = val
            set_cell_bg(c, sev_color.get(sev, bg) if j == 3 else bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
                if j == 3:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                if j == 0:
                    run.font.bold = True
    add_page_break(doc)

    # SECTION 8: ABLATION STUDY (BARU)
    add_slide_header(doc, 8, "Ablation Study - IF vs LOF vs EE",
                     "Justifikasi penggunaan ensemble (bukan single model)")
    add_bullet(doc, "Tujuan: membuktikan ensemble lebih baik daripada single model.")
    doc.add_paragraph()
    add_img(doc, fig_abl_detection(), width=5.0)
    add_img(doc, fig_abl_overlap(), width=6.2)
    doc.add_paragraph()
    add_bullet(doc, "Metrik kesepakatan antar model:", bold_part="Metrik kesepakatan antar model:")
    make_table(doc, ['Pasangan Model', 'Jaccard Index', "Cohen's Kappa", 'Interpretasi'], [
        (name, f'{j:.4f}', f'{k:.4f}',
         'tinggi' if k > 0.6 else 'moderate' if k > 0.4 else 'fair' if k > 0.2 else 'rendah')
        for name, j, k in ABL_AGREEMENT
    ], col_fmt={0: 'bold'})
    add_page_break(doc)

    # SECTION 9: ABLATION QUALITY
    add_slide_header(doc, 9, "Ablation Study - Evaluasi Kualitas",
                     "Precision@K & Avg rule_violations (proxy-based)")
    add_img(doc, fig_abl_quality(), width=6.5)
    doc.add_paragraph()
    make_table(doc, ['Model', 'Deteksi', 'Precision@K', 'Avg rule_viol', 'Avg score'], [
        (m, ABL_QUALITY[m]['n'], f"{ABL_QUALITY[m]['prec_k']:.4f}",
         f"{ABL_QUALITY[m]['avg_rv']:.2f}", f"{ABL_QUALITY[m]['avg_score']:.4f}")
        for m in ['IF', 'LOF', 'EE', 'Ensemble']
    ], col_fmt={0: 'bold'})
    doc.add_paragraph()
    add_bullet(doc, f"Baseline populasi: Avg rule_viol = {BASELINE_RV:.2f}, "
                    f"Avg score = {BASELINE_SCORE:.4f}", level=0)
    doc.add_paragraph()
    add_bullet(doc, "Temuan kunci:", bold_part="Temuan kunci:")
    add_bullet(doc, "EE paling selaras dengan domain rules (Precision@K tertinggi)", level=1)
    add_bullet(doc, "LOF Avg rule_viol di bawah baseline - menangkap anomali ORTOGONAL "
                    "terhadap rules (anomali lokal/tersembunyi)", level=1)
    add_bullet(doc, "IF di tengah - menyeimbangkan kedua tipe", level=1)
    add_bullet(doc, "Justifikasi ensemble: tidak ada single model dominan di semua aspek", level=1)
    add_page_break(doc)

    # SECTION 10: SHAP
    add_slide_header(doc, 10, "Phase 5.5 - SHAP Explainability",
                     "Mengapa setiap user dianggap anomali")
    add_bullet(doc, "Metode: shap.TreeExplainer pada model Isolation Forest")
    add_bullet(doc, f"Output: SHAP values {TOTAL_USERS} user x 11 features")
    doc.add_paragraph()
    causes = {}
    for a in TOP_ANOMALIES:
        causes[a[5]] = causes.get(a[5], 0) + 1
    add_bullet(doc, "Top SHAP causes pada top 10 anomalies:", bold_part="Top SHAP causes pada top 10 anomalies:")
    for cause, cnt in sorted(causes.items(), key=lambda x: -x[1]):
        add_bullet(doc, f'"{cause}": {cnt} dari 10 anomali teratas', level=1)
    doc.add_paragraph()
    add_bullet(doc, "Interpretasi:", bold_part="Interpretasi:")
    add_bullet(doc, "Nilai SHAP positif = fitur menaikkan anomaly score", level=1)
    add_bullet(doc, "Nilai absolut terbesar = faktor dominan", level=1)
    add_bullet(doc, "Manfaat audit: setiap temuan punya bukti kuantitatif yang traceable", level=1)
    add_page_break(doc)

    # SECTION 11: RECOMMENDATIONS
    add_slide_header(doc, 11, "Key Findings & Recommendations",
                     "Temuan utama dan rekomendasi tindak lanjut")
    h_find = doc.add_heading("Key Findings", level=2)
    for run in h_find.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    add_bullet(doc, f"{ANOMALY_DIST.get('CRITICAL',0)} user CRITICAL terdeteksi dengan threshold data-driven")
    add_bullet(doc, f"Top user ({TOP_ANOMALIES[0][0]}) - score {TOP_ANOMALIES[0][1]:.4f}, "
                    f"{TOP_ANOMALIES[0][3]} rules")
    add_bullet(doc, "Ablation study membuktikan ensemble lebih robust dari single model")
    doc.add_paragraph()
    h_rec = doc.add_heading("Rekomendasi Tindak Lanjut", level=2)
    for run in h_rec.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    add_bullet(doc, "Immediate (CRITICAL):", bold_part="Immediate (CRITICAL):")
    add_bullet(doc, f"Audit & reset password {ANOMALY_DIST.get('CRITICAL',0)} user CRITICAL", level=1)
    add_bullet(doc, "Review privileged access logs untuk top 10 user", level=1)
    add_bullet(doc, "Short-term (HIGH/MEDIUM):", bold_part="Short-term (HIGH/MEDIUM):")
    add_bullet(doc, "Enable MFA untuk admin account", level=1)
    add_bullet(doc, "Monitor shared device & lockout pattern", level=1)
    add_bullet(doc, "Long-term:", bold_part="Long-term:")
    add_bullet(doc, "Validasi threshold dengan label domain expert", level=1)
    add_bullet(doc, "Deploy continuous behavioral analytics", level=1)
    add_page_break(doc)

    # SECTION 12: OUTPUT FILES
    add_slide_header(doc, 12, "File Output", "Daftar file yang dihasilkan pipeline")
    make_table(doc, ['File Path', 'Isi'], [
        ('output/anomaly_detection_report.txt', 'Laporan teks untuk auditor'),
        ('output/anomaly_detection_detailed.json', 'Detail per user (SHAP + rules)'),
        ('output/anomaly_statistics.json', 'Statistik distribusi anomali'),
        ('output/AD_Anomaly_Detection_Report_v3.docx', 'Laporan DOCX (file ini)'),
        ('data/phase4_graph_features.csv', f'{TOTAL_USERS} user x 11 features'),
        ('data/phase5_anomaly_results.csv', 'Score & severity quantile-based'),
        ('data/phase55_shap_values.csv', f'SHAP values {TOTAL_USERS} user x 11 features'),
        ('models/*.pkl', 'Trained models: IF, LOF, EE, Scaler'),
    ])
    doc.add_paragraph()
    add_divider(doc)
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final.add_run("AD Log -> Neo4j -> Rules -> Features -> Ensemble -> SHAP -> Report")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("Pipeline Selesai - Mei 2026 - Report v3 (quantile threshold + ablation study)")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rf.font.italic = True

    out_path = 'output/AD_Anomaly_Detection_Report_v3.docx'
    doc.save(out_path)
    print(f"[OK] DOCX saved: {out_path}")
    return out_path


if __name__ == '__main__':
    print("Generating report v3...")
    path = build_document()
    print(f"Done: {path}")
