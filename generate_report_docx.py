#!/usr/bin/env python3
"""
Generate Project Report DOCX
AD Anomaly Detection - Graph-Based Knowledge System
15 Slide Sections
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import io
import os

os.makedirs('output', exist_ok=True)

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_slide_header(doc, slide_num, title, subtitle=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ SLIDE {slide_num} ]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.bold = False

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(20)

    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(11)
            run.font.italic = True

    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('─' * 70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
    doc.add_paragraph()

def add_bullet(doc, text, level=0, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(text[len(bold_part):])
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

def add_img(doc, fig, width=6.0):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    plt.close(fig)
    doc.add_paragraph()

def add_page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────
# IMAGE GENERATORS
# ─────────────────────────────────────────────

def fig_pipeline():
    fig, ax = plt.subplots(figsize=(12, 3))
    ax.axis('off')
    stages = [
        ("AD Log\nData", "#1A3C6E"),
        ("Neo4j\nKnowledge\nGraph", "#1565C0"),
        ("Rule-Based\nKnowledge\nEngine", "#1976D2"),
        ("Graph\nFeature\nExtraction", "#1E88E5"),
        ("Isolation\nForest\n+ Ensemble", "#42A5F5"),
        ("SHAP\nExplainability", "#64B5F6"),
        ("Anomaly\nDetection\nOutput", "#90CAF9"),
    ]
    n = len(stages)
    w, h_box, gap = 1.4, 0.7, 0.15
    total = n * w + (n - 1) * gap
    x0 = (12 - total) / 2

    for i, (label, color) in enumerate(stages):
        x = x0 + i * (w + gap)
        rect = FancyBboxPatch((x, 0.15), w, h_box, boxstyle="round,pad=0.05",
                               facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w / 2, 0.15 + h_box / 2, label, ha='center', va='center',
                fontsize=7.5, color='white', fontweight='bold', linespacing=1.4)
        if i < n - 1:
            ax.annotate('', xy=(x + w + gap, 0.15 + h_box / 2),
                        xytext=(x + w, 0.15 + h_box / 2),
                        arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        ax.text(x + w / 2, 0.08, f"Phase {i+1 if i < 5 else ('5.5' if i == 5 else '6')}",
                ha='center', va='center', fontsize=6.5, color='#666')

    ax.set_xlim(0, 12)
    ax.set_ylim(0, 1.1)
    fig.tight_layout(pad=0)
    return fig

def fig_graph_schema():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.axis('off')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    nodes = {
        'User':       (5.0, 5.0, '#1A3C6E'),
        'Hostname':   (2.0, 3.5, '#1565C0'),
        'Server':     (8.0, 3.5, '#1976D2'),
        'IPAddress':  (2.0, 1.5, '#1E88E5'),
        'Group':      (8.0, 1.5, '#42A5F5'),
        'Event':      (5.0, 2.5, '#64B5F6'),
        'Service':    (5.0, 0.5, '#90CAF9'),
    }

    edges = [
        ('User', 'Hostname',  'LOGIN_FROM',     'left'),
        ('User', 'Server',    'AUTHENTICATED_VIA', 'right'),
        ('User', 'IPAddress', 'CONNECTED_FROM', 'left'),
        ('User', 'Group',     'MEMBER_OF',      'right'),
        ('Event','User',      'REFERENCES',     'down'),
        ('Hostname','IPAddress','USED_IP',       'down'),
        ('User', 'Service',   'USED_SERVICE',   'down'),
    ]

    for label, (x, y, color) in nodes.items():
        circle = plt.Circle((x, y), 0.45, color=color, zorder=3)
        ax.add_patch(circle)
        ax.text(x, y, label, ha='center', va='center', fontsize=7,
                color='white', fontweight='bold', zorder=4, wrap=True)

    for src, dst, rel, _ in edges:
        x1, y1, _ = nodes[src]
        x2, y2, _ = nodes[dst]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#888', lw=1.5,
                                   connectionstyle='arc3,rad=0.1'), zorder=2)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my, rel, ha='center', va='center', fontsize=6,
                color='#333', style='italic',
                bbox=dict(boxstyle='round,pad=0.2', fc='#F0F4FF', ec='none'))

    ax.set_title('Neo4j Knowledge Graph — Node & Relationship Schema', fontsize=11,
                 fontweight='bold', color='#1A3C6E', pad=10)
    fig.tight_layout()
    return fig

def fig_rules():
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis('off')
    rules = [
        ('R001', 'Normal Login Hosts',        'Unique hosts > 3'),
        ('R002', 'Business Hours Pattern',     'Off-hours logins > 10%'),
        ('R003', 'Shared Device Detection',    'Device used by > 5 users'),
        ('R004', 'Uncommon Server Access',     'Critical server accessed'),
        ('R005', 'Failed Login Spike',         'Failures > 50 total'),
        ('R006', 'Unusual IP Address',         'Non-office / non-VPN IP'),
        ('R007', 'After-Hours Privileged',     'Admin + off-hours + critical'),
    ]
    cols = ['Rule ID', 'Rule Name', 'Trigger Condition']
    col_w = [0.10, 0.45, 0.45]
    y_start = 0.92
    row_h = 0.115

    for j, (col, cw) in enumerate(zip(cols, col_w)):
        x = sum(col_w[:j])
        rect = FancyBboxPatch((x + 0.005, y_start), cw - 0.01, row_h,
                               boxstyle="round,pad=0.01", facecolor='#1A3C6E', edgecolor='none')
        ax.add_patch(rect)
        ax.text(x + cw / 2, y_start + row_h / 2, col, ha='center', va='center',
                color='white', fontsize=8.5, fontweight='bold',
                transform=ax.transAxes)

    for i, (rid, name, cond) in enumerate(rules):
        y = y_start - (i + 1) * row_h
        bg = '#EEF4FF' if i % 2 == 0 else '#F9FAFF'
        for j, (text, cw) in enumerate(zip([rid, name, cond], col_w)):
            x = sum(col_w[:j])
            rect = FancyBboxPatch((x + 0.005, y), cw - 0.01, row_h,
                                   boxstyle="round,pad=0.01", facecolor=bg, edgecolor='none')
            ax.add_patch(rect)
            clr = '#1A3C6E' if j == 0 else '#333'
            fw = 'bold' if j == 0 else 'normal'
            ax.text(x + cw / 2, y + row_h / 2, text, ha='center', va='center',
                    color=clr, fontsize=8, fontweight=fw,
                    transform=ax.transAxes)

    ax.set_title('Rule-Based Knowledge Engine — 7 Domain Rules', fontsize=11,
                 fontweight='bold', color='#1A3C6E', pad=8)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    return fig

def fig_features():
    features = [
        'Host Diversity', 'Critical Server\nAccess Ratio', 'Failed Login\nRatio',
        'Shared Device\nRisk', 'IP Network\nRisk', 'Privilege\nLevel',
        'Graph\nConnectivity', 'Rule\nViolations'
    ]
    importance = [0.18, 0.14, 0.22, 0.10, 0.16, 0.08, 0.06, 0.06]
    colors = ['#1A3C6E', '#1565C0', '#1976D2', '#1E88E5',
              '#2196F3', '#42A5F5', '#64B5F6', '#90CAF9']

    fig, ax = plt.subplots(figsize=(10, 3.5))
    bars = ax.barh(features, importance, color=colors, edgecolor='white', linewidth=0.8)
    ax.set_xlabel('Relative Contribution to Anomaly Detection', fontsize=9)
    ax.set_title('8 Graph-Based Features — Indicative Importance', fontsize=11,
                 fontweight='bold', color='#1A3C6E')
    ax.set_xlim(0, 0.28)
    for bar, val in zip(bars, importance):
        ax.text(val + 0.003, bar.get_y() + bar.get_height() / 2,
                f'{val:.0%}', va='center', fontsize=8, color='#333')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return fig

def fig_ensemble():
    fig, axes = plt.subplots(1, 3, figsize=(10, 3.5))
    methods = ['Isolation Forest\n(IF)', 'Local Outlier\nFactor (LOF)', 'Elliptic\nEnvelope (EE)']
    colors  = ['#1A3C6E', '#1976D2', '#42A5F5']
    descs   = [
        'Random partitioning\nTree-based\nFast & scalable',
        'Density-based\nLocal neighborhoods\nNon-parametric',
        'Robust covariance\nGaussian assumption\nStatistical'
    ]
    for ax, method, color, desc in zip(axes, methods, colors, descs):
        circle = plt.Circle((0.5, 0.55), 0.35, color=color, transform=ax.transAxes,
                             clip_on=False, zorder=2)
        ax.add_patch(circle)
        ax.text(0.5, 0.55, method, ha='center', va='center', fontsize=8.5,
                color='white', fontweight='bold', transform=ax.transAxes, zorder=3)
        ax.text(0.5, 0.12, desc, ha='center', va='center', fontsize=7.5,
                color='#555', transform=ax.transAxes, linespacing=1.5)
        ax.axis('off')

    fig.suptitle('Ensemble Anomaly Detection — 3 Independent Methods\n'
                 'Final Score = 60% Ensemble Votes + 40% Rule Violations',
                 fontsize=10, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_shap():
    features = ['failure_ratio', 'ip_network_risk', 'host_diversity',
                 'rule_violations', 'shared_device_risk', 'privilege_level',
                 'critical_server_ratio', 'connectivity']
    shap_vals = [0.31, 0.24, 0.18, 0.12, 0.08, 0.05, 0.02, -0.01]
    colors = ['#C62828' if v >= 0 else '#1565C0' for v in shap_vals]

    fig, ax = plt.subplots(figsize=(9, 3.8))
    bars = ax.barh(features, shap_vals, color=colors, edgecolor='white', linewidth=0.8)
    ax.axvline(0, color='#999', linewidth=0.8, linestyle='--')
    ax.set_xlabel('SHAP Value (contribution to anomaly score)', fontsize=9)
    ax.set_title('SHAP Explainability — Example: User john.doe (Anomaly Score: 0.87)',
                 fontsize=10, fontweight='bold', color='#1A3C6E')
    for bar, val in zip(bars, shap_vals):
        offset = 0.005 if val >= 0 else -0.005
        ha = 'left' if val >= 0 else 'right'
        ax.text(val + offset, bar.get_y() + bar.get_height() / 2,
                f'{val:+.2f}', va='center', ha=ha, fontsize=8, color='#333')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    red_patch  = mpatches.Patch(color='#C62828', label='Increases anomaly score')
    blue_patch = mpatches.Patch(color='#1565C0', label='Decreases anomaly score')
    ax.legend(handles=[red_patch, blue_patch], fontsize=8, loc='lower right')
    fig.tight_layout()
    return fig

def fig_severity():
    labels = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'NORMAL']
    sizes  = [5, 12, 18, 15, 50]
    colors = ['#B71C1C', '#E53935', '#FB8C00', '#FDD835', '#43A047']
    explode = (0.08, 0.05, 0, 0, 0)

    fig, ax = plt.subplots(figsize=(7, 4))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, explode=explode,
        autopct='%1.0f%%', startangle=140, pctdistance=0.75,
        textprops={'fontsize': 9}
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title('Expected Anomaly Severity Distribution\n(714 Users Analyzed)',
                 fontsize=10, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_execution_status():
    phases = [
        'Phase 1: Data Preparation',
        'Phase 2: Neo4j Ingestion',
        'Phase 3: Rule-Based Engine',
        'Phase 4: Feature Extraction',
        'Phase 5: Anomaly Detection',
        'Phase 5.5: SHAP',
        'Phase 6: Reporting',
    ]
    statuses = ['Done', 'Running', 'Pending', 'Pending', 'Pending', 'Pending', 'Pending']
    colors_map = {'Done': '#43A047', 'Running': '#FB8C00', 'Pending': '#BDBDBD'}
    bar_colors = [colors_map[s] for s in statuses]
    progress   = [100, 60, 0, 0, 0, 0, 0]

    fig, ax = plt.subplots(figsize=(9, 3.5))
    bars = ax.barh(phases, progress, color=bar_colors, edgecolor='white',
                   linewidth=0.8, height=0.6)
    ax.set_xlim(0, 120)
    ax.set_xlabel('Progress (%)', fontsize=9)
    ax.set_title('Pipeline Execution Status', fontsize=11, fontweight='bold', color='#1A3C6E')
    for bar, status, val in zip(bars, statuses, progress):
        ax.text(val + 1, bar.get_y() + bar.get_height() / 2,
                f'{status}', va='center', fontsize=8.5, color='#333')
    ax.tick_params(labelsize=8)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    patches = [mpatches.Patch(color=v, label=k) for k, v in colors_map.items()]
    ax.legend(handles=patches, fontsize=8, loc='lower right')
    fig.tight_layout()
    return fig

# ─────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── SLIDE 1: COVER ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN TEKNIS")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    title_p = doc.add_heading('Deteksi Anomali Active Directory\nBerbasis Graph Knowledge System', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        run.font.size = Pt(24)

    doc.add_paragraph()
    sub = doc.add_paragraph('AD Log → Neo4j Knowledge Graph → Rule-Based Engine →\nGraph Feature Extraction → Isolation Forest → SHAP → Anomaly Output')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(11)
        run.font.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph('April 2026  ·  7 Phases  ·  1,833,352 Events  ·  714 Users')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.5)

    add_page_break(doc)

    # ── SLIDE 2: PROBLEM STATEMENT ──────────────────────────────────
    add_slide_header(doc, 2, "Problem Statement",
                     "Mengapa pendekatan lama tidak cukup?")

    add_bullet(doc, "Pendekatan lama (problematic):", bold_part="Pendekatan lama (problematic):")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.6)
    r = p.add_run("AD Log  →  Behavioral Features  →  Isolation Forest  →  Neo4j (Visualisasi saja)")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    r.font.bold = True

    doc.add_paragraph()
    add_bullet(doc, "Masalah dengan pendekatan lama:")
    add_bullet(doc, "Neo4j hanya dipakai untuk visualisasi, bukan engine kecerdasan", level=1)
    add_bullet(doc, "Isolation Forest langsung dari raw features — tidak ada context relasi", level=1)
    add_bullet(doc, "Anomali terdeteksi tapi tidak bisa dijelaskan (black-box)", level=1)
    add_bullet(doc, "Tidak scalable untuk query relasi multi-hop", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Solusi yang diusulkan:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.6)
    r = p.add_run("AD Log  →  Neo4j (Core Intelligence)  →  Rules  →  Features  →  IF  →  SHAP  →  Output")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    r.font.bold = True

    doc.add_paragraph()
    add_bullet(doc, "Neo4j sebagai knowledge engine utama, bukan storage/visualisasi")
    add_bullet(doc, "Domain rules ditanamkan di graph sebelum ML berjalan")
    add_bullet(doc, "SHAP menjelaskan KENAPA setiap user dianggap anomali")

    add_divider(doc)
    add_page_break(doc)

    # ── SLIDE 3: ARCHITECTURE OVERVIEW ──────────────────────────────
    add_slide_header(doc, 3, "Architecture Overview",
                     "7-Phase Pipeline — AD Log ke Anomaly Output")

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Phase', 'Nama', 'Output Utama']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_bg(cell, '1A3C6E')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(10)

    rows_data = [
        ('Phase 1', 'Data Preparation', '1.8M events, 12 kolom, CSV bersih'),
        ('Phase 2', 'Neo4j Knowledge Graph', '714 User, 7 node types, 8 relationship types'),
        ('Phase 3', 'Rule-Based Knowledge Engine', '7 rules dievaluasi per user, violation flags'),
        ('Phase 4', 'Graph Feature Extraction', '8 fitur per user diekstrak dari graph'),
        ('Phase 5', 'Ensemble Anomaly Detection', 'IF + LOF + EE → anomaly score + severity'),
        ('Phase 5.5', 'SHAP Explainability', 'SHAP values per feature per anomalous user'),
        ('Phase 6', 'Reporting', 'Text report + JSON + statistics'),
    ]
    for i, (phase, name, output) in enumerate(rows_data):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        for j, text in enumerate([phase, name, output]):
            cell = table.cell(i + 1, j)
            cell.text = text
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if j == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    add_img(doc, fig_pipeline(), width=6.2)
    add_page_break(doc)

    # ── SLIDE 4: AD LOG DATA ─────────────────────────────────────────
    add_slide_header(doc, 4, "Phase 1 — AD Log Data",
                     "Sumber data: unified_logon_events.csv")

    add_bullet(doc, "Statistik dataset:", bold_part="Statistik dataset:")
    add_bullet(doc, "Total events: 1,833,352 baris", level=1)
    add_bullet(doc, "Total users: 714 unique users", level=1)
    add_bullet(doc, "Total hostnames: 1,273 unique devices", level=1)
    add_bullet(doc, "Total IP addresses: 1,275 unique IPs", level=1)
    add_bullet(doc, "Servers: 12 (termasuk Domain Controllers)", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Kolom data (12 fields):", bold_part="Kolom data (12 fields):")

    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    cols_data = [
        ('event_source', 'Sumber event (Windows Security, Sysmon, etc.)'),
        ('username', 'Username AD'),
        ('hostname', 'Nama device/komputer'),
        ('ip_address', 'Alamat IP sumber'),
        ('timestamp', 'Waktu event (ISO format)'),
        ('success', 'Status login (True/False)'),
        ('event_type + failure_reason + dc_name + server_name + domain + logon_service', '6 kolom tambahan konteks'),
    ]
    for i, (col, desc) in enumerate(cols_data):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        c1 = table2.cell(i, 0)
        c2 = table2.cell(i, 1)
        c1.text = col
        c2.text = desc
        for c in [c1, c2]:
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
        for run in c1.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_page_break(doc)

    # ── SLIDE 5: NEO4J SCHEMA ────────────────────────────────────────
    add_slide_header(doc, 5, "Phase 2 — Neo4j Knowledge Graph",
                     "Representasi relasi AD sebagai property graph")

    add_bullet(doc, "7 Node Types:")
    for n in ['User', 'Hostname', 'Server', 'IPAddress', 'Group', 'Service', 'Event']:
        add_bullet(doc, n, level=1)

    doc.add_paragraph()
    add_bullet(doc, "8 Relationship Types:")
    rels = ['LOGIN_FROM (User→Hostname)', 'AUTHENTICATED_VIA (User→Server)',
            'FAILED_LOGIN (User→Server)', 'CONNECTED_FROM (User→IPAddress)',
            'USED_IP (Hostname→IPAddress)', 'USED_SERVICE (User→Service)',
            'REFERENCES (Event→User)', 'MEMBER_OF (User→Group)']
    for r in rels:
        add_bullet(doc, r, level=1)

    doc.add_paragraph()
    add_img(doc, fig_graph_schema(), width=6.0)
    add_page_break(doc)

    # ── SLIDE 6: GRAPH EXAMPLE ───────────────────────────────────────
    add_slide_header(doc, 6, "Phase 2 — Contoh Graph Traversal",
                     "Query relasi multi-hop yang tidak bisa dilakukan SQL")

    add_bullet(doc, "Contoh query Cypher — temukan user anomali dengan multi-hop:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    code = (
        "MATCH (u:User)-[:LOGIN_FROM]->(h:Hostname)\n"
        "MATCH (u)-[:CONNECTED_FROM]->(ip:IPAddress)\n"
        "MATCH (u)-[:FAILED_LOGIN]->(s:Server)\n"
        "WHERE ip.range_category NOT IN ['Office_Network','VPN']\n"
        "AND u.rule_R005_total_failures > 50\n"
        "RETURN u.username, count(DISTINCT h), count(DISTINCT ip)"
    )
    r = p.add_run(code)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    add_bullet(doc, "Kenapa tidak bisa pakai SQL/RDBMS:")
    add_bullet(doc, "Query di atas butuh 3–4 JOIN di SQL — lambat pada 1.8M rows", level=1)
    add_bullet(doc, "Neo4j traversal = O(1) per relationship, tidak bergantung total data", level=1)
    add_bullet(doc, "Graph algorithms (centrality, path) tersedia native di Neo4j", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Hasil ingestion Phase 2:")

    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    stats = [('Total nodes', '714 + 1.273 + 12 + 1.275 + 7 + 3 + 1.833.352 = ~1.8M nodes'),
             ('Total relationships', '~10M relationships (8 types)'),
             ('Constraints & Indexes', '7 unique constraints + 7 indexes')]
    for i, (k, v) in enumerate(stats):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        c1 = table3.cell(i + 1, 0)
        c2 = table3.cell(i + 1, 1)
        c1.text = k
        c2.text = v
        for c in [c1, c2]:
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
        for run in c1.paragraphs[0].runs:
            run.font.bold = True
    hrow = table3.row_cells(0)
    for j, h in enumerate(['Metric', 'Value']):
        hrow[j].text = h
        set_cell_bg(hrow[j], '1A3C6E')
        for run in hrow[j].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    add_page_break(doc)

    # ── SLIDE 7: RULE-BASED ENGINE ───────────────────────────────────
    add_slide_header(doc, 7, "Phase 3 — Rule-Based Knowledge Engine",
                     "7 domain rules dievaluasi langsung di Neo4j via Cypher")

    add_img(doc, fig_rules(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Output per user setelah Phase 3:")
    add_bullet(doc, "u.rule_R001_violation = true/false", level=1)
    add_bullet(doc, "u.rule_violations = total jumlah rules yang dilanggar (0–7)", level=1)
    add_bullet(doc, "u.max_rule_severity = HIGH / MEDIUM / LOW", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Rule violations dijadikan Feature 8 untuk input Isolation Forest")

    add_page_break(doc)

    # ── SLIDE 8: FEATURES ────────────────────────────────────────────
    add_slide_header(doc, 8, "Phase 4 — Graph Feature Extraction",
                     "8 fitur diekstrak dari relasi Neo4j — bukan dari kolom CSV langsung")

    add_img(doc, fig_features(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Kenapa features dari graph, bukan langsung dari CSV?")
    add_bullet(doc, "CSV hanya menyimpan event individual — tidak ada konteks relasi", level=1)
    add_bullet(doc, "Graph menghitung HUBUNGAN: berapa device, server, IP unik per user?", level=1)
    add_bullet(doc, "Feature 7 (Graph Connectivity) = degree centrality — tidak bisa dari CSV", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Output: data/phase4_graph_features.csv — 714 rows × 8 kolom, siap input IF")

    add_page_break(doc)

    # ── SLIDE 9: ANOMALY DETECTION ───────────────────────────────────
    add_slide_header(doc, 9, "Phase 5 — Ensemble Anomaly Detection",
                     "3 metode independen — ensemble voting untuk akurasi lebih tinggi")

    add_img(doc, fig_ensemble(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Formula final anomaly score:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("Final Score = (Ensemble Votes / 3) × 0.60  +  (Rule Violations / 7) × 0.40")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    add_bullet(doc, "User dianggap anomali jika:")
    add_bullet(doc, "2+ metode setuju = anomali, ATAU", level=1)
    add_bullet(doc, "Final score > 0.75", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Severity classification:")
    add_bullet(doc, "CRITICAL: score ≥ 0.75", level=1)
    add_bullet(doc, "HIGH: score 0.60–0.74", level=1)
    add_bullet(doc, "MEDIUM: score 0.45–0.59", level=1)
    add_bullet(doc, "LOW: score 0.30–0.44", level=1)
    add_bullet(doc, "NORMAL: score < 0.30", level=1)

    add_page_break(doc)

    # ── SLIDE 10: SHAP ───────────────────────────────────────────────
    add_slide_header(doc, 10, "Phase 5.5 — SHAP Explainability",
                     "Transformasi black-box IF menjadi penjelasan per-feature per-user")

    add_img(doc, fig_shap(), width=6.0)

    doc.add_paragraph()
    add_bullet(doc, "Metode: shap.TreeExplainer — native support untuk Isolation Forest")
    add_bullet(doc, "Input: Trained IF model + 8 graph features × 714 users")
    add_bullet(doc, "Output: SHAP value per fitur per user anomali")

    doc.add_paragraph()
    add_bullet(doc, "Interpretasi SHAP:")
    add_bullet(doc, "Nilai positif = fitur ini menaikkan anomaly score", level=1)
    add_bullet(doc, "Nilai negatif = fitur ini menurunkan anomaly score", level=1)
    add_bullet(doc, "Nilai absolut terbesar = faktor paling dominan", level=1)

    doc.add_paragraph()
    add_bullet(doc, "Kenapa SHAP penting untuk audit?")
    add_bullet(doc, "Auditor bisa jelaskan kepada manajemen: 'User X anomali karena failure_ratio 95%'", level=1)
    add_bullet(doc, "Bukan sekedar label 'anomali' tanpa bukti kuantitatif", level=1)

    add_page_break(doc)

    # ── SLIDE 11: REPORTING ──────────────────────────────────────────
    add_slide_header(doc, 11, "Phase 6 — Reporting & Output",
                     "3 format output: human-readable, detailed JSON, summary statistics")

    add_bullet(doc, "File output yang dihasilkan:")
    outputs = [
        ('anomaly_detection_report.txt', 'Laporan teks untuk auditor/manajemen'),
        ('anomaly_detection_detailed.json', 'Bukti detail per user (SHAP + rules + graph)'),
        ('anomaly_statistics.json', 'Statistik distribusi anomali'),
        ('phase5_anomaly_results.csv', 'Score + severity semua 714 user'),
        ('models/', 'Model IF, LOF, EE, Scaler tersimpan'),
    ]
    table4 = doc.add_table(rows=len(outputs) + 1, cols=2)
    table4.style = 'Table Grid'
    for j, h in enumerate(['File', 'Isi']):
        c = table4.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)
    for i, (fname, desc) in enumerate(outputs):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        c1 = table4.cell(i + 1, 0)
        c2 = table4.cell(i + 1, 1)
        c1.text = fname
        c2.text = desc
        for c in [c1, c2]:
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
        for run in c1.paragraphs[0].runs:
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    add_img(doc, fig_severity(), width=5.5)
    add_page_break(doc)

    # ── SLIDE 12: WHY NOT RDBMS / EXCEL / IF-ELSE ───────────────────
    add_slide_header(doc, 12, "Justifikasi Arsitektur",
                     "Kenapa bukan RDBMS, Excel, atau if-else saja?")

    h2 = doc.add_heading("vs. RDBMS (SQL)", level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    add_bullet(doc, "Query relasi multi-hop (User→Hostname→IP) butuh 3–4 JOIN di SQL")
    add_bullet(doc, "Neo4j: traversal O(1) per relationship, tidak bergantung jumlah data")
    add_bullet(doc, "Graph algorithms (degree centrality, path) tidak ada native di SQL")
    add_bullet(doc, "Relasi many-to-many seperti login AD sangat mahal di RDBMS")

    doc.add_paragraph()
    h3 = doc.add_heading("vs. Excel", level=2)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    add_bullet(doc, "Excel tidak bisa handle 1.8 juta baris secara efisien")
    add_bullet(doc, "Tidak ada konsep graph traversal atau relationship query")
    add_bullet(doc, "Tidak ada audit trail yang queryable")
    add_bullet(doc, "Tidak scalable untuk monitoring AD real-time")

    doc.add_paragraph()
    h4 = doc.add_heading("vs. Pure If-Else (Rule Only)", level=2)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    add_bullet(doc, "If-else hanya mendeteksi anomali yang SUDAH diketahui (predefined patterns)")
    add_bullet(doc, "Isolation Forest mendeteksi pola TIDAK TERDUGA yang tidak masuk rule manapun")
    add_bullet(doc, "Kombinasi Rules + IF memberikan dua lapisan: konteks domain + anomali statistik")
    add_bullet(doc, "Contoh: user normal = 15 failures, user jahat = 5 — if-else gagal, IF tetap deteksi")

    add_page_break(doc)

    # ── SLIDE 13: DATA FLOW EXAMPLE ──────────────────────────────────
    add_slide_header(doc, 13, "End-to-End Data Flow",
                     "Contoh: Deteksi user anomali step-by-step")

    steps = [
        ("Step 1 — Raw Event",
         "User: john.doe  |  Host: PC-UNUSUAL  |  Server: DC01\n"
         "Timestamp: 02:30 AM  |  IP: 10.50.x.x  |  Failed: 95x"),
        ("Step 2 — Graph Ingestion (Phase 2)",
         "Node User, Hostname, Server, IPAddress dibuat/update\n"
         "8 relationship types ditambahkan ke graph"),
        ("Step 3 — Rule Evaluation (Phase 3)",
         "R002 VIOLATION: login jam 02:30 (off-hours)\n"
         "R005 VIOLATION: 95 failed logins\n"
         "R006 VIOLATION: IP 10.50.x.x bukan office/VPN\n"
         "rule_violations = 3"),
        ("Step 4 — Feature Extraction (Phase 4)",
         "failure_ratio = 0.95  |  ip_network_risk = 1.0\n"
         "host_diversity = 2.3  |  rule_violations = 3"),
        ("Step 5 — Isolation Forest (Phase 5)",
         "IF score: -0.78 (anomali)\n"
         "LOF: anomali  |  EE: anomali  →  3/3 votes\n"
         "Final Score: 0.87  |  Severity: CRITICAL"),
        ("Step 6 — SHAP (Phase 5.5)",
         "failure_ratio: +0.31 (faktor terbesar)\n"
         "ip_network_risk: +0.24  |  rule_violations: +0.12"),
        ("Step 7 — Report Output (Phase 6)",
         "User john.doe → CRITICAL\n"
         "Alasan: failure ratio 95%, IP tidak biasa, login dini hari"),
    ]

    for i, (title, content) in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"  {title}  ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(4)
        # Fake highlight with arrow prefix
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.5)
        r2 = p2.add_run(content)
        r2.font.size = Pt(9.5)
        r2.font.name = 'Courier New'
        r2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        if i < len(steps) - 1:
            arr = doc.add_paragraph("    ↓")
            arr.paragraph_format.left_indent = Inches(0.3)
            for run in arr.runs:
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.size = Pt(10)

    add_page_break(doc)

    # ── SLIDE 14: EXECUTION STATUS ───────────────────────────────────
    add_slide_header(doc, 14, "Execution Status",
                     "Progress pipeline per 30 April 2026")

    add_img(doc, fig_execution_status(), width=6.0)

    doc.add_paragraph()
    status_data = [
        ('Phase 1', 'Data Preparation', 'DONE', '1.8M events, 12 kolom, validated'),
        ('Phase 2', 'Neo4j Ingestion', 'RUNNING', 'Fix timestamp malformed (1-10-08 00:00:00)'),
        ('Phase 3', 'Rule-Based Engine', 'PENDING', 'Ready — 7 rules Cypher siap'),
        ('Phase 4', 'Feature Extraction', 'PENDING', 'Ready — 8 features siap'),
        ('Phase 5', 'Anomaly Detection', 'PENDING', 'IF + LOF + EE ensemble'),
        ('Phase 5.5', 'SHAP', 'PENDING', 'Script belum dibuat — setelah Phase 5'),
        ('Phase 6', 'Reporting', 'PENDING', 'Perlu update include SHAP values'),
    ]
    t5 = doc.add_table(rows=len(status_data) + 1, cols=4)
    t5.style = 'Table Grid'
    for j, h in enumerate(['Phase', 'Nama', 'Status', 'Catatan']):
        c = t5.cell(0, j)
        c.text = h
        set_cell_bg(c, '1A3C6E')
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)
    status_colors = {'DONE': '43A047', 'RUNNING': 'FB8C00', 'PENDING': 'BDBDBD'}
    for i, (ph, name, status, note) in enumerate(status_data):
        bg = 'EEF4FF' if i % 2 == 0 else 'F9FAFF'
        for j, text in enumerate([ph, name, status, note]):
            c = t5.cell(i + 1, j)
            c.text = text
            set_cell_bg(c, bg if j != 2 else status_colors.get(status, 'BDBDBD'))
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                if j == 2:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True

    add_page_break(doc)

    # ── SLIDE 15: OUTCOMES & NEXT STEPS ─────────────────────────────
    add_slide_header(doc, 15, "Expected Outcomes & Next Steps",
                     "Target akhir pipeline dan rekomendasi tindak lanjut")

    h_out = doc.add_heading("Expected Outcomes", level=2)
    for run in h_out.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_bullet(doc, "714 user dianalisis dengan 8 graph features + 7 rule evaluations")
    add_bullet(doc, "Anomaly score 0–1 per user dengan label severity (CRITICAL/HIGH/MEDIUM/LOW/NORMAL)")
    add_bullet(doc, "SHAP values mengidentifikasi faktor dominan per anomali")
    add_bullet(doc, "3 output report: text, JSON detail, statistik")
    add_bullet(doc, "Setiap temuan memiliki audit trail yang traceable")

    doc.add_paragraph()
    h_next = doc.add_heading("Next Steps Setelah Pipeline Selesai", level=2)
    for run in h_next.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_bullet(doc, "Immediate:", bold_part="Immediate:")
    add_bullet(doc, "Tunggu Phase 2 selesai (Neo4j ingestion 1.8M events)", level=1)
    add_bullet(doc, "Jalankan Phase 3 → 4 → 5 secara berurutan", level=1)

    add_bullet(doc, "Setelah baseline selesai:", bold_part="Setelah baseline selesai:")
    add_bullet(doc, "Buat script neo4j_phase55_shap.py (SHAP explainability)", level=1)
    add_bullet(doc, "Update Phase 6 reporting untuk include SHAP values", level=1)

    add_bullet(doc, "Enhancement opsional:", bold_part="Enhancement opsional:")
    add_bullet(doc, "Node2Vec + clustering untuk graph-native anomaly detection", level=1)
    add_bullet(doc, "LSTM Autoencoder untuk temporal sequence anomaly", level=1)
    add_bullet(doc, "GenAI layer untuk narasi root cause analysis otomatis", level=1)

    doc.add_paragraph()
    add_divider(doc)

    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final.add_run("AD Log → Neo4j → Rules → Features → Isolation Forest → SHAP → Anomaly Output")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("Dokumen ini dibuat otomatis oleh pipeline generator  ·  April 2026")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rf.font.italic = True

    out_path = 'output/AD_Anomaly_Detection_Report.docx'
    doc.save(out_path)
    print(f"[OK] DOCX saved: {out_path}")
    return out_path


if __name__ == '__main__':
    print("Generating report...")
    path = build_document()
    print(f"Done: {path}")
