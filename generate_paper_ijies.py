#!/usr/bin/env python3
"""
Generate IJIES Draft Paper
AD Anomaly Detection - Graph-Based Knowledge System

Format: International Journal of Intelligent Engineering and Systems (IJIES)
Two-column, Times New Roman, structured for journal submission.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pandas as pd
import json
import os
import re

os.makedirs('output', exist_ok=True)

# ── LOAD DATA ────────────────────────────────────────────────────────
with open('output/anomaly_statistics.json') as f:
    STATS = json.load(f)

TOTAL_USERS  = STATS['total_users']
ANOMALY_DIST = STATS['anomaly_distribution']
SCORE_STATS  = STATS['anomaly_score_stats']
ENS_VOTE     = STATS['ensemble_voting']
RULE_STATS   = STATS['rule_violations_stats']

df = pd.read_csv('data/phase5_anomaly_results.csv').drop_duplicates(subset='user_id').reset_index(drop=True)
shap_df = pd.read_csv('data/phase55_shap_values.csv')

P75 = df['final_anomaly_score'].quantile(0.75)
P90 = df['final_anomaly_score'].quantile(0.90)
P95 = df['final_anomaly_score'].quantile(0.95)
P99 = df['final_anomaly_score'].quantile(0.99)

# Top 5 anomalies with SHAP
# The pipeline stores SHAP "top cause" labels in Indonesian (Indonesian-facing
# dashboards/explainer); translate them to English for the journal paper only.
# .get(x, x) passes through any value already in English (e.g. after a re-run).
EN_LABEL = {
    'Login dari banyak host':    'Logins from many hosts',
    'Akses critical server':     'Critical-server access',
    'Intensitas login gagal':    'Failed-login intensity',
    'Risiko shared device':      'Shared-device risk',
    'Risiko IP tidak dikenal':   'Unknown-IP risk',
    'Level privilege tinggi':    'High privilege level',
    'Graph connectivity tinggi': 'High graph connectivity',
    'Pelanggaran rule':          'Rule violations',
    'Sering lockout':            'Frequent lockouts',
    'Banyak admin action':       'Excessive admin actions',
    'Anggota grup sensitif':     'Sensitive-group membership',
}
merged = df.merge(shap_df[['user_id', 'top_feature_1_label']], on='user_id', how='left')
merged['top_feature_1_label'] = merged['top_feature_1_label'].map(
    lambda x: EN_LABEL.get(x, x) if pd.notna(x) else x)
TOP5 = merged.nlargest(5, 'final_anomaly_score')

TOTAL_ANOMALI = int(df['severity'].isin(['CRITICAL', 'HIGH', 'MEDIUM']).sum())

# ── COMBINATION ABLATION (rank-based fusion vs rule-based proxy) ──────
from scipy.stats import rankdata as _rankdata
_MR = {
    'IF':  _rankdata(df['if_score'].values),
    'LOF': _rankdata(df['lof_score'].values),
    'EE':  _rankdata(df['ee_score'].values),
}
_NN = len(df)
_GT = (df['rule_violations'].values >= 6).astype(int)
_PP = int(_GT.sum())
_CONFIGS = [('IF',), ('LOF',), ('EE',), ('IF', 'LOF'),
            ('IF', 'EE'), ('LOF', 'EE'), ('IF', 'LOF', 'EE')]

def _eval_cfg(models, gt, K):
    combined = sum(_MR[m] for m in models) / len(models)
    order = combined.argsort()
    pred = [0] * _NN
    for idx in order[:K]:
        pred[idx] = 1
    pred = pd.Series(pred).values
    TP = int(((pred == 1) & (gt == 1)).sum()); FP = int(((pred == 1) & (gt == 0)).sum())
    FN = int(((pred == 0) & (gt == 1)).sum()); TN = int(((pred == 0) & (gt == 0)).sum())
    acc = (TP + TN) / _NN
    prec = TP / (TP + FP) if TP + FP else 0
    rec = TP / (TP + FN) if TP + FN else 0
    f1 = 2 * prec * rec / (prec + rec) if prec + rec else 0
    return acc, prec, rec, f1

COMBO_ABLATION = []
for c in _CONFIGS:
    acc, prec, rec, f1 = _eval_cfg(c, _GT, _PP)
    COMBO_ABLATION.append(('+'.join(c), acc, prec, rec, f1))

# Sensitivity: F1 of each config across proxy thresholds; thresholds chosen to
# remain statistically meaningful (each has a non-trivial positive count).
_SENS_T = [t for t in [4, 5, 6] if int((df['rule_violations'].values >= t).sum()) >= 20]
SENS_TABLE = {}   # config -> {t: f1}
SENS_BEST = {}    # t -> (config, f1)
for t in _SENS_T:
    gt_t = (df['rule_violations'].values >= t).astype(int)
    Kt = int(gt_t.sum())
    scored = [('+'.join(c), _eval_cfg(c, gt_t, Kt)[3]) for c in _CONFIGS]
    for name, f1v in scored:
        SENS_TABLE.setdefault(name, {})[t] = f1v
    SENS_BEST[t] = max(scored, key=lambda x: x[1])

# Is the full ensemble ever the worst configuration? (robustness claim)
_ENS_NEVER_WORST = all(
    SENS_TABLE['IF+LOF+EE'][t] > min(SENS_TABLE[c][t] for c in SENS_TABLE)
    for t in _SENS_T
)
# Does a single config dominate (win) across all thresholds?
_SENS_WINNERS = {SENS_BEST[t][0] for t in _SENS_T}

# ── INDIVIDUAL-MODEL ABLATION (Precision@K vs rule proxy) ────────────
def _model_set(col):
    return set(df[df[col] == 1]['user_id'])
_HEAVY = set(df[df['rule_violations'] >= 6]['user_id'])
_BASE_RV = df['rule_violations'].mean()

IND_ABLATION = []  # (name, detected, precision@k, avg_rule_viol, avg_score)
for name, col in [('IF', 'if_anomaly'), ('LOF', 'lof_anomaly'),
                  ('EE', 'ee_anomaly'), ('Ensemble', None)]:
    s = set(df[df['anomaly_votes'] >= 2]['user_id']) if col is None else _model_set(col)
    sub = df[df['user_id'].isin(s)]
    n = len(s)
    pk = len(s & _HEAVY) / n if n else 0
    IND_ABLATION.append((name, n, pk, sub['rule_violations'].mean() if n else 0,
                         sub['final_anomaly_score'].mean() if n else 0))

# ── INTER-MODEL AGREEMENT (Jaccard, Cohen's Kappa) ───────────────────
from sklearn.metrics import cohen_kappa_score as _kappa
def _jac(a, b):
    return len(a & b) / len(a | b) if (a | b) else 0.0
_S = {'IF': _model_set('if_anomaly'), 'LOF': _model_set('lof_anomaly'), 'EE': _model_set('ee_anomaly')}
AGREEMENT = []  # (pair, jaccard, kappa, interp)
for a, b, ca, cb in [('IF', 'LOF', 'if_anomaly', 'lof_anomaly'),
                     ('IF', 'EE', 'if_anomaly', 'ee_anomaly'),
                     ('LOF', 'EE', 'lof_anomaly', 'ee_anomaly')]:
    j = _jac(_S[a], _S[b]); k = _kappa(df[ca], df[cb])
    interp = ('High' if k > 0.6 else 'Moderate' if k > 0.4 else 'Fair' if k > 0.2 else 'Low')
    AGREEMENT.append((f'{a}–{b}', j, k, interp))

# Convenience handles used across abstract + results sections
_ee = next(r for r in IND_ABLATION if r[0] == 'EE')
_lof = next(r for r in IND_ABLATION if r[0] == 'LOF')
_lofee = next(a for a in AGREEMENT if a[0] == 'LOF–EE')

# ── DEPLOYED-FLAG CONFUSION MATRIX vs rule-based proxy ───────────────
_cm_pred = ((df['anomaly_votes'] >= 2) | (df['final_anomaly_score'] > 0.75)).astype(int).values
_cm_act = (df['rule_violations'] >= 6).astype(int).values
_CM_TP = int(((_cm_pred == 1) & (_cm_act == 1)).sum())
_CM_FP = int(((_cm_pred == 1) & (_cm_act == 0)).sum())
_CM_FN = int(((_cm_pred == 0) & (_cm_act == 1)).sum())
_CM_TN = int(((_cm_pred == 0) & (_cm_act == 0)).sum())
_CM_P = _CM_TP / (_CM_TP + _CM_FP) if (_CM_TP + _CM_FP) else 0.0
_CM_R = _CM_TP / (_CM_TP + _CM_FN) if (_CM_TP + _CM_FN) else 0.0
_CM_F1 = 2 * _CM_P * _CM_R / (_CM_P + _CM_R) if (_CM_P + _CM_R) else 0.0
_CM_ACC = (_CM_TP + _CM_TN) / len(df)

# ── CONFUSION-MATRIX HEATMAP (Figure 3) — generated from the same counts ──
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as _plt
    import numpy as _np
    _cm = _np.array([[_CM_TP, _CM_FN], [_CM_FP, _CM_TN]])
    _fig, _ax = _plt.subplots(figsize=(4.4, 3.8))
    _ax.imshow(_cm, cmap='Blues')
    _ax.set_xticks([0, 1]); _ax.set_xticklabels(['Flagged', 'Not flagged'])
    _ax.set_yticks([0, 1]); _ax.set_yticklabels(['Proxy\npositive', 'Proxy\nnegative'])
    _ax.set_xlabel('Predicted'); _ax.set_ylabel('Actual (rule-based proxy)')
    _lab = [['TP', 'FN'], ['FP', 'TN']]
    _thr = _cm.max() / 2.0
    for _i in range(2):
        for _j in range(2):
            _ax.text(_j, _i, f'{_lab[_i][_j]}\n{_cm[_i, _j]:,}', ha='center', va='center',
                     fontsize=12, fontweight='bold',
                     color='white' if _cm[_i, _j] > _thr else '#111111')
    _fig.tight_layout()
    os.makedirs('docs/diagrams', exist_ok=True)
    _fig.savefig('docs/diagrams/confusion_matrix.png', dpi=200, bbox_inches='tight')
    _plt.close(_fig)
except Exception as _e:
    print(f'[warn] confusion-matrix figure not generated: {_e}')

# ── SHAP FEATURE-IMPORTANCE BAR CHART (Figure 2) — mean |SHAP| per feature ──
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as _plt
    _shap_cols = [c for c in shap_df.columns if c.startswith('shap_')]
    _imp = shap_df[_shap_cols].abs().mean().sort_values()  # ascending -> largest on top
    _names = [c.replace('shap_', '') for c in _imp.index]
    _fig, _ax = _plt.subplots(figsize=(6.2, 3.6))
    _colors = ['#9aa7b8' if v == 0 else '#2c5f8a' for v in _imp.values]
    _ax.barh(_names, _imp.values, color=_colors)
    _ax.set_xlabel('Mean |SHAP value|')
    _xmax = max(_imp.values) if len(_imp) else 1.0
    for _i, _v in enumerate(_imp.values):
        _ax.text(_v + _xmax * 0.01, _i, f'{_v:.3f}', va='center', fontsize=8)
    _ax.margins(x=0.14)
    _ax.spines['top'].set_visible(False)
    _ax.spines['right'].set_visible(False)
    _ax.tick_params(labelsize=8)
    _fig.tight_layout()
    os.makedirs('docs/diagrams', exist_ok=True)
    _fig.savefig('docs/diagrams/shap_importance.png', dpi=200, bbox_inches='tight')
    _plt.close(_fig)
except Exception as _e:
    print(f'[warn] SHAP importance figure not generated: {_e}')

# ── HELPERS ──────────────────────────────────────────────────────────

def set_two_columns(section):
    """Set section to two-column layout"""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols[0].set(qn('w:num'), '2')
        cols[0].set(qn('w:space'), '720')   # 0.5 inch space between columns
    else:
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')
        sectPr.append(cols)

def add_section_break(doc, columns=2):
    """Add section break and set columns"""
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols[0].set(qn('w:num'), str(columns))
        cols[0].set(qn('w:space'), '720')
    else:
        cols_el = OxmlElement('w:cols')
        cols_el.set(qn('w:num'), str(columns))
        cols_el.set(qn('w:space'), '720')
        sectPr.append(cols_el)
    return new_section

def set_font(run, name='Times New Roman', size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def add_hyperlink(paragraph, url, text, size=11, color='0563C1'):
    """Append a clickable external hyperlink run (blue, underlined) to a paragraph."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_internal_link(paragraph, anchor, text, size=11):
    """Append a clickable run that jumps to an internal bookmark (citation cross-ref).
    Kept as normal black text (journal style) — clickable but not visually colored."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rPr.append(col)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

_CITE_GROUP_RE = re.compile(r'\[\s*\d+(?:\s*,\s*\d+)*\s*\]')   # [7], [2, 12]
_NUM_RE = re.compile(r'\d+')

def add_cited_runs(paragraph, text, size=11):
    """Add text, turning bracketed citation numbers ([7], [2, 12], and the
    endpoints of ranges like [16]-[19]) into clickable internal links to the
    matching reference bookmark (ref_N). Brackets/commas/dashes stay as text."""
    pos = 0
    for mg in _CITE_GROUP_RE.finditer(text):
        if mg.start() > pos:
            r = paragraph.add_run(text[pos:mg.start()]); set_font(r, size=size)
        group = mg.group(0)
        gpos = 0
        for nm in _NUM_RE.finditer(group):
            if nm.start() > gpos:
                r = paragraph.add_run(group[gpos:nm.start()]); set_font(r, size=size)
            add_internal_link(paragraph, f'ref_{nm.group(0)}', nm.group(0), size=size)
            gpos = nm.end()
        if gpos < len(group):
            r = paragraph.add_run(group[gpos:]); set_font(r, size=size)
        pos = mg.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:]); set_font(r, size=size)

def add_bookmark_to_paragraph(paragraph, name, bm_id):
    """Place an (empty) bookmark at the start of a paragraph — a jump target for
    in-text citation cross-references."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bm_id)); start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), str(bm_id))
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(start)
    else:
        paragraph._p.insert(0, start)
    start.addnext(end)
    return name

def add_para(doc, text, size=11, bold=False, italic=False, align='left', indent=True, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading_lvl1(doc, num, text):
    """1. First-order heading - 12pt bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f'{num}. {text}')
    set_font(run, size=12, bold=True)
    return p

def add_heading_lvl2(doc, num, text):
    """1.1 Second-order heading - 11pt bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f'{num} {text}')
    set_font(run, size=11, bold=True)
    return p

def add_text(doc, text, justify=True, indent=True, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    add_cited_runs(p, text, size=11)
    return p

def add_equation(doc, eq, num):
    """Indented equation with number"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(eq)
    set_font(run, size=11, italic=True)
    # Equation number right-aligned via tab
    tab_run = p.add_run(f'\t({num})')
    set_font(tab_run, size=11)
    # Set right tab
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.0), WD_ALIGN_PARAGRAPH.RIGHT)

def add_figure(doc, img_path, caption_num, caption, width_cm=17.0, span_columns=True):
    """Full-width figure (spanning both columns) with IJIES caption below (9pt centered)."""
    if span_columns:
        add_section_break(doc, columns=1)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(width_cm))
        pic = doc.paragraphs[-1]
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(6)
        pic.paragraph_format.space_after = Pt(2)
    else:
        miss = doc.add_paragraph()
        miss.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = miss.add_run(f'[Figure missing: {img_path}]')
        set_font(mr, size=9, italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    crun = cap.add_run(f'Figure {caption_num}. {caption}')
    set_font(crun, size=9)
    if span_columns:
        add_section_break(doc, columns=2)

def _set_tight_cell_margins(table, lr_twips=40, tb_twips=20):
    """Shrink left/right cell margins so dense tables fit inside one journal
    column without spanning. Default Word margin is 108 twips per side; 40 twips
    (~0.7 mm) reclaims enough room to fit 5–6 columns at 8pt."""
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(el)
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('left', lr_twips), ('right', lr_twips),
                      ('top', tb_twips), ('bottom', tb_twips)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)

def add_table(doc, headers, rows, caption_num, caption, font_size=8):
    """Table with IJIES caption above (10pt centered) and a compact 8pt body.

    Tables stay WITHIN the two-column layout (no column spanning, so the
    template is preserved). An 8pt body plus tightened cell margins lets even
    dense 5–6 column tables fit a single journal column, with Word autofit
    sizing each column to its content so words wrap at spaces, never mid-word.
    """
    n = len(headers)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    crun = cap.add_run(f'Table {caption_num}. {caption}')
    set_font(crun, size=9)
    # Table
    table = doc.add_table(rows=len(rows) + 1, cols=n)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_tight_cell_margins(table)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        run = c.paragraphs[0].add_run(str(h))
        set_font(run, size=font_size, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(i + 1, j)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            set_font(run, size=font_size)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add empty line after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

def add_schema_box(doc, pairs, col1_cm=3.0, col2_cm=4.8, font_size=8):
    """Compact, UN-NUMBERED 2-column reference box (label | members), left-aligned.

    Presents dense identifier lists (e.g. node/relationship types) cleanly,
    without stretching justified prose into whitespace rivers and without
    disturbing the numbered-table (Table 1..N) sequence — it carries no caption.
    """
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    _set_tight_cell_margins(table)
    for i, (label, members) in enumerate(pairs):
        c0 = table.cell(i, 0); c0.text = ''
        r0 = c0.paragraphs[0].add_run(label)
        set_font(r0, size=font_size, bold=True)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        c1 = table.cell(i, 1); c1.text = ''
        r1 = c1.paragraphs[0].add_run(members)
        set_font(r1, size=font_size)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.cell(i, 0).width = Cm(col1_cm)
        table.cell(i, 1).width = Cm(col2_cm)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table

# ── BUILD PAPER ──────────────────────────────────────────────────────

def build_paper():
    doc = Document()

    # Page setup: A4, margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)

    # ── TITLE BLOCK (single column) ──────────────────────────────────
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('Explainable Anomaly Detection in Active Directory: Integrating a '
                    'Rule-Based Knowledge Engine, Ensemble Learning, and SHAP for '
                    'Human-Readable Reasoning')
    set_font(run, size=14, bold=True)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Mahathir Muhammad')
    set_font(run, size=11, bold=True)
    sup = p.add_run('1')
    set_font(sup, size=11, bold=True)
    sup.font.superscript = True
    run = p.add_run('      ')
    set_font(run, size=11, bold=True)
    run = p.add_run('Kelly Rossa Sungkono')
    set_font(run, size=11, bold=True)
    sup = p.add_run('2*')
    set_font(sup, size=11, bold=True)
    sup.font.superscript = True
    run = p.add_run('      ')
    set_font(run, size=11, bold=True)
    run = p.add_run('Riyanarto Sarno')
    set_font(run, size=11, bold=True)
    sup = p.add_run('3')
    set_font(sup, size=11, bold=True)
    sup.font.superscript = True

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    sup = p.add_run('1,2,3')
    set_font(sup, size=11)
    sup.font.superscript = True
    run = p.add_run(' Institut Teknologi Sepuluh Nopember, Indonesia')
    set_font(run, size=11, italic=True)

    # Email
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('* Corresponding author’s Email: kelly@its.ac.id')
    set_font(run, size=10, italic=True)

    # Horizontal line (simulated)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('_' * 100)
    set_font(run, size=8)

    # Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run('Abstract: ')
    set_font(run, size=10, bold=True)
    run = p.add_run(
        'We present a graph-based pipeline that both detects and explains anomalies in '
        'Active Directory (AD) authentication logs. AD produces millions of events a day, and '
        'the rule-based or statistical tools normally used to sift them either lose the '
        'surrounding context or hand analysts a verdict with no reasoning behind it. Our '
        'design pushes the logs into a Neo4j knowledge graph and scores users with three '
        'deliberately dissimilar unsupervised models: Isolation Forest (IF), Local Outlier '
        'Factor (LOF), and Elliptic Envelope (EE). Lacking any ground truth, we fix severity '
        'from the data through quantile thresholds (P75/P90/P95/P99). SHAP TreeExplainer '
        'attributes each user’s score, and a knowledge-base-grounded explainer (Phase 7) '
        'rewrites the result as plain-language reasoning that cites MITRE ATT&CK techniques '
        'and Windows Event IDs, with a deterministic template as fallback when no LLM is '
        'available. '
        f'Over {TOTAL_USERS:,} users built from 1.8 million logon events, the pipeline flags '
        f'{TOTAL_ANOMALI} at MEDIUM severity or above: '
        f'CRITICAL ({ANOMALY_DIST.get("CRITICAL",0)}), HIGH ({ANOMALY_DIST.get("HIGH",0)}), '
        f'and MEDIUM ({ANOMALY_DIST.get("MEDIUM",0)}), the CRITICAL and HIGH portion coming '
        'from the ensemble majority vote. Across all seven configurations the models stay '
        f'complementary: EE hugs the rule engine (Precision@K={_ee[2]:.1%}), LOF finds '
        'orthogonal local anomalies, and IF sits in between. Their weak agreement '
        f'(Cohen’s Kappa LOF–EE={_lofee[2]:.2f}), and the full ensemble never finishing last '
        'across proxy thresholds, justifies using all three. We also show the flattering '
        'imbalanced accuracy (~0.93) to be an artifact of majority-class prediction, since F1 '
        'jumps from 0.46 to 0.78 once classes are balanced; F1, not accuracy, is therefore '
        'our headline figure. Throughout, SHAP keeps the per-user explanations auditable for '
        'security analysts.')
    set_font(run, size=10)

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run('Keywords: ')
    set_font(run, size=10, bold=True)
    run = p.add_run('Anomaly detection, Knowledge graph, Active Directory, Ensemble learning, '
                    'Isolation Forest, SHAP, Explainable AI, Cybersecurity.')
    set_font(run, size=10)

    # ── SWITCH TO TWO COLUMN ─────────────────────────────────────────
    add_section_break(doc, columns=2)

    # ── 1. INTRODUCTION ──────────────────────────────────────────────
    add_heading_lvl1(doc, 1, 'Introduction')
    add_text(doc,
        'Active Directory (AD) sits at the center of identity and access management for most '
        'organizations. The problem is volume: even one domain records millions of '
        'authentication events a day, so hand-review is hopeless. Teams rely on automated '
        'anomaly detection instead, to catch compromised accounts, insider activity, or '
        'policy violations early [2, 12].')
    add_text(doc,
        'Three limitations recur in existing AD anomaly detection. Rule-based systems are '
        'easy to read, yet they only fire on patterns someone has already encoded, so '
        'genuinely new behavior slips through. Classical machine-learning detectors treat '
        'each event in isolation and discard the relational structure AD is full of (who '
        'logged in from which device, onto which server, over which connection). And many '
        'ML systems behave as black boxes: they raise an alert but cannot tell an analyst '
        'why [13, 14].')
    add_text(doc,
        'Our pipeline tackles all three at once. It combines a knowledge-graph view of the '
        'data, a domain rule engine, a heterogeneous unsupervised ensemble, and SHAP-based '
        'explanations. Concretely, this work contributes the following:')
    add_text(doc,
        '(1) We map AD audit logs onto a Neo4j knowledge graph of seven node types and ten '
        'relationship types, which opens up traversal queries that relational tables cannot '
        'answer efficiently.')
    add_text(doc,
        '(2) We combine three unsupervised models (IF, LOF, EE) by majority vote, betting '
        'that their differing inductive biases will catch different kinds of anomaly.')
    add_text(doc,
        '(3) Severity bands come from the data, not from us: we cut them at the '
        'P75/P90/P95/P99 quantiles, following recent anomaly-detection work [1, 3].')
    add_text(doc,
        '(4) Every flagged user carries a SHAP TreeExplainer breakdown, so an opaque score '
        'becomes something an analyst can actually act on.')
    add_text(doc,
        '(5) An ablation weighs each ensemble member’s contribution with the Jaccard index, '
        'Cohen’s Kappa, and Precision@K, all against a rule-based proxy.')

    # ── 2. RELATED WORK ──────────────────────────────────────────────
    add_heading_lvl1(doc, 2, 'Related work')
    add_text(doc,
        'Anomaly detection over security and authentication logs is still a moving target. '
        'Tufan et al. [2] bring machine learning to anomaly-based intrusion detection on an '
        'institutional network, while Marteau [1] introduces a random-partitioning forest (an '
        'Isolation-Forest cousin) for point-wise and collective anomalies; Xu et al. [3] later '
        'add learned representations on top in Deep Isolation Forest. When the input is log '
        'data, Le and Zhang [5] skip log parsing entirely, and Landauer et al. [4] give a '
        'broad survey of deep-learning log-anomaly methods. Recent IEEE work pushes ensemble '
        'and deep-learning intrusion detection further, through stacking, hybrid, and '
        'broad-learning ensembles [16]–[19], [21], explainable stacking ensembles [20], and '
        'detectors built for robustness or spatial-temporal signals [30], [31].')
    add_text(doc,
        'Graph-based representations have gained traction because they make multi-hop '
        'relations explicit (for instance, user → device → IP), relations that are costly to '
        'reconstruct with relational joins. Sun and Yang [7] cast lateral movement as '
        'anomalous links and detect it with a heterogeneous graph neural network, whereas '
        'Smiliotopoulos et al. [6] take a supervised route over Windows Sysmon logs. On the '
        'advanced-persistent-threat side, Jia et al. [8] learn masked graph representations '
        'over provenance graphs, and Zhou et al. [9] pair learned models with logical '
        'reasoning on graph data, an idea close in spirit to the knowledge-driven part of our '
        'own pipeline. GNNs have since seen wide use for anomaly and intrusion detection '
        '[22]–[26]: federated variants [22], dynamic line-graph models [23], spectral GNNs '
        '[24], and insider-threat detectors [26].')
    add_text(doc,
        'User-behavior and insider-threat analytics sit right next to AD anomaly detection. '
        'Xiao et al. [10] build multi-timescale interaction graphs of user behavior for '
        'insider-threat detection; Roy and Chen [11] fold cyber-human factors into a deep '
        'framework; and Alzaabi and Mehmood [12] survey where machine learning currently '
        'stands on malicious insider detection.')
    add_text(doc,
        'Explanation is no longer optional in a security operations center. Gaspar et al. '
        '[13] put LIME and SHAP onto intrusion-detection models, Barnard et al. [14] lean on '
        'explainable AI to harden a detector, and Nascita et al. [15] survey the field for '
        'traffic classification and intrusion detection. Useful as they are, these efforts '
        'mostly stop at numeric attributions. The same explainable-AI toolkit has also '
        'reached smart-grid, IoT, and high-dimensional cybersecurity problems [27]–[29].')
    add_text(doc,
        'Our work differs by unifying, within a single Active Directory pipeline, a rule-based '
        'knowledge engine, a heterogeneous unsupervised ensemble (IF + LOF + EE), and SHAP '
        'attribution, and — beyond prior work — by converting the result into human-readable '
        'explanations grounded in a security knowledge base with verbatim MITRE ATT&CK and '
        'Windows Event ID citations (Phase 7).')

    # ── 3. METHODOLOGY ───────────────────────────────────────────────
    add_heading_lvl1(doc, 3, 'Proposed method')

    add_heading_lvl2(doc, '3.1', 'System architecture')
    add_text(doc,
        'The proposed pipeline consists of eight sequential phases: (1) Data preparation, '
        '(2) Neo4j knowledge graph ingestion, (3) Rule-based engine, (4) Graph feature '
        'extraction, (5) Ensemble anomaly detection, (5.5) SHAP explainability, (6) '
        'Reporting, and (7) knowledge-base-grounded human-readable explanation. The pipeline '
        'transforms raw AD log events into actionable anomaly reports with human-readable, '
        'source-attributed explanations. The overall architecture is shown in Fig. 1.')

    add_figure(doc, 'docs/diagrams/framework_figure_paper.png', 1,
        'Overview of the proposed explainable anomaly-detection framework. AD audit logs '
        'first become a Neo4j knowledge graph, which a rule-based engine and graph feature '
        'extraction then process; a heterogeneous unsupervised ensemble (Isolation Forest, '
        'Local Outlier Factor, Elliptic Envelope) scores the result. Flagged users are ranked '
        'by data-driven quantile severity and attributed with SHAP, after which a '
        'generative-AI (Gen-AI) explainer, grounded in a security knowledge base (MITRE '
        'ATT&CK and Windows Event IDs), turns each case into a human-readable explanation '
        '(Phase 7) ahead of evaluation.')

    add_heading_lvl2(doc, '3.2', 'Knowledge graph construction')
    add_text(doc,
        'We load the AD audit logs into Neo4j as a property graph whose schema, listed '
        'below, has seven node types and ten relationship types:')
    add_schema_box(doc, [
        ('Node types (7)',
         'User, Hostname, Server, IPAddress, Service, Group, Event'),
        ('Relationship types (10)',
         'LOGIN_FROM, AUTHENTICATED_VIA, FAILED_LOGIN, CONNECTED_FROM, USED_IP, '
         'USED_SERVICE, MEMBER_OF, REFERENCES, LOCKED_OUT, ADMIN_ACTION_ON'),
    ])
    add_text(doc,
        'LOCKED_OUT and ADMIN_ACTION_ON feed straight into the lockout_count and '
        'admin_actions features. A single log row may spawn several edges, each carrying its '
        'timestamp and event metadata as properties.')

    add_heading_lvl2(doc, '3.3', 'Rule-based knowledge engine')
    add_text(doc,
        'Expert knowledge enters the pipeline as ten domain rules, written in Cypher and run '
        'straight against the graph. They look for tell-tale patterns such as logging in from '
        'many hosts (R001), activity outside working hours (R002), bursts of failed logins '
        '(R005), or an unusual amount of admin activity (R009). The number of rules a user '
        'trips, rule_violations, runs from 0 to 10; we store it on the node both as a signal '
        'analysts can read and as a feature the later ML stages consume.')

    add_heading_lvl2(doc, '3.4', 'Graph feature extraction')
    add_text(doc,
        'From the graph we derive eleven per-user features: host diversity, critical server '
        'access ratio, failed-login intensity, shared device risk, IP network risk, privilege '
        'level, graph connectivity (degree centrality), rule violations, lockout count, admin '
        'actions, and sensitive group membership. Failed-login intensity deserves a note: it '
        'divides the summed failed-login counts by the number of login relationships, and '
        'because the two quantities carry different units the result is unbounded (it can run '
        'into the thousands) rather than a 0–1 proportion. None of these features fall out of '
        'row-level CSV data without traversing the graph, which is exactly why the graph '
        'representation earns its place.')

    add_heading_lvl2(doc, '3.5', 'Ensemble anomaly detection')
    add_text(doc,
        'On the standardized feature matrix we fit three unsupervised models, each with a '
        'contamination rate of 0.05: Isolation Forest (tree-based), Local Outlier Factor '
        '(density-based), and Elliptic Envelope (statistical). Every model votes on each user, '
        'and we label a user anomalous when at least two of the three agree, or when the final '
        'fused score climbs above 0.75.')
    add_text(doc, 'The final anomaly score combines voting and rule violations:')
    add_equation(doc,
        'final_score = 0.60 × (votes / 3) + 0.40 × (rule_violations / 10)', 1)

    add_heading_lvl2(doc, '3.6', 'Quantile-based severity classification')
    add_text(doc,
        'We avoid fixing severity cut-offs by fiat and instead read them off the data with '
        'quantile thresholds. The score distribution is heavily skewed '
        f'(min={SCORE_STATS["min"]:.3f}, max={SCORE_STATS["max"]:.3f}, '
        f'mean={SCORE_STATS["mean"]:.3f}), so a flat absolute cut-off, say a direct CVSS '
        'mapping, would dump nearly everyone into a single class and leave the others almost '
        'empty.')
    add_text(doc, 'The five-tier classification follows:')
    add_equation(doc,
        f'CRITICAL: score ≥ P99 ({P99:.4f})', 2)
    add_equation(doc,
        f'HIGH: P95 ({P95:.4f}) ≤ score < P99', 3)
    add_equation(doc,
        f'MEDIUM: P90 ({P90:.4f}) ≤ score < P95', 4)
    add_equation(doc,
        f'LOW: P75 ({P75:.4f}) ≤ score < P90', 5)
    add_text(doc,
        'The choice of P95 specifically aligns with the contamination rate (5%) of the '
        'three ML models, ensuring consistency between model assumptions and classification '
        'thresholds.')

    add_heading_lvl2(doc, '3.7', 'SHAP explainability')
    add_text(doc,
        'Per-user explanations are generated using SHAP TreeExplainer on the trained '
        'Isolation Forest model. For each user, SHAP values quantify the contribution of '
        'each feature to the anomaly score. The top contributing feature is recorded as the '
        '"top cause" label, enabling security analysts to understand why a user was flagged '
        'without inspecting raw scores.')

    add_heading_lvl2(doc, '3.8', 'Human-readable explanation (knowledge-based)')
    add_text(doc,
        'Raw anomaly scores and SHAP attributions are still difficult for analysts to act on '
        'directly. Phase 7 — the Gen-AI explainer shown in Fig. 1 — converts each flagged user '
        'into a human-readable explanation '
        'grounded in a curated security knowledge base (KB) that maps the eleven features and '
        'ten rules to verified MITRE ATT&CK techniques, Windows Event IDs, and recommended '
        'mitigations. Two explanation modes share the same KB: a deterministic template that '
        'always runs offline with zero hallucination, and an optional LLM-generated narrative '
        '(GPT-4o-mini, temperature 0.2, strict JSON schema) that is strictly grounded to the '
        'KB facts supplied in the prompt. Citations (technique IDs/URLs and Event IDs) are '
        'taken verbatim from the KB rather than produced by the model, and the system falls '
        'back to the template when no API key is available or the model output fails '
        'validation. This makes the final explanations simultaneously human-readable, '
        'reproducible, and source-attributed.')

    # ── 4. EXPERIMENTAL SETUP ────────────────────────────────────────
    add_heading_lvl1(doc, 4, 'Experimental setup')

    add_heading_lvl2(doc, '4.1', 'Dataset')
    add_text(doc,
        'Our data is a 1,833,352-event export of AD logons drawn from several domain '
        'controllers and member servers, produced by ManageEngine ADAudit Plus. Once '
        f'ingested into the graph it resolves to {TOTAL_USERS:,} distinct users, 1,273 '
        'hostnames, 1,275 IP addresses, and 7 servers, each user tied to their login events, '
        'authentication outcomes, group memberships, and admin activity.')

    add_heading_lvl2(doc, '4.2', 'Implementation')
    add_text(doc,
        'The pipeline is implemented in Python 3.12, with Neo4j 5.x as the graph database. '
        'Machine learning models use scikit-learn 1.4 with default parameters except for '
        'contamination=0.05. SHAP values are computed using shap 0.46 TreeExplainer. All '
        'experiments were conducted on a standard workstation (Windows 11, 16GB RAM).')

    add_heading_lvl2(doc, '4.3', 'Evaluation methodology')
    add_text(doc,
        'Unsupervised anomaly detection hands us no ground-truth labels, so we fall back on a '
        'proxy: any user with rule_violations ≥ 6 is taken as a "true anomaly." Against that '
        'proxy we score every ensemble configuration with the usual metrics (Accuracy, '
        'Precision, Recall, F1), plus Precision@K, the Jaccard index, and Cohen’s Kappa for '
        'agreement between models. The caveat matters: since the proxy comes from the rule '
        'engine itself, these numbers are weak supervision and will favor detectors that '
        'already think like the rules. For that reason we sweep several proxy thresholds '
        'rather than trusting any single one.')

    # ── 5. RESULTS ───────────────────────────────────────────────────
    add_heading_lvl1(doc, 5, 'Results and discussion')

    add_heading_lvl2(doc, '5.1', 'Severity distribution')
    add_text(doc,
        f'The quantile-based severity split is laid out in Table 1. In total '
        f'{TOTAL_ANOMALI} users ({100*TOTAL_ANOMALI/TOTAL_USERS:.1f}%) land at MEDIUM '
        f'severity or above, and {ANOMALY_DIST.get("CRITICAL",0)} of them reach CRITICAL '
        f'(the top 1%).')

    add_table(doc,
        ['Severity', 'Threshold', 'Users', 'Percentage'],
        [
            ('CRITICAL', f'≥ {P99:.4f}', ANOMALY_DIST.get('CRITICAL', 0),
             f'{100*ANOMALY_DIST.get("CRITICAL",0)/TOTAL_USERS:.1f}%'),
            ('HIGH',     f'≥ {P95:.4f}', ANOMALY_DIST.get('HIGH', 0),
             f'{100*ANOMALY_DIST.get("HIGH",0)/TOTAL_USERS:.1f}%'),
            ('MEDIUM',   f'≥ {P90:.4f}', ANOMALY_DIST.get('MEDIUM', 0),
             f'{100*ANOMALY_DIST.get("MEDIUM",0)/TOTAL_USERS:.1f}%'),
            ('LOW',      f'≥ {P75:.4f}', ANOMALY_DIST.get('LOW', 0),
             f'{100*ANOMALY_DIST.get("LOW",0)/TOTAL_USERS:.1f}%'),
            ('NORMAL',   f'< {P75:.4f}',      ANOMALY_DIST.get('NORMAL', 0),
             f'{100*ANOMALY_DIST.get("NORMAL",0)/TOTAL_USERS:.1f}%'),
        ], 1, 'Severity distribution with quantile-based thresholds')

    add_heading_lvl2(doc, '5.2', 'Top anomalous users')
    add_text(doc,
        'Table 2 presents the top five anomalous users ranked by anomaly score, with their '
        'rule violation counts, ensemble voting, and SHAP-derived top causes. The highest-'
        f'scoring user ({TOP5.iloc[0]["username"]}) achieves a score of '
        f'{TOP5.iloc[0]["final_anomaly_score"]:.4f} with all three models in agreement and '
        f'{int(TOP5.iloc[0]["rule_violations"])} out of 10 rules violated.')

    top_rows = []
    for _, r in TOP5.iterrows():
        top_rows.append((
            r['username'],
            f'{r["final_anomaly_score"]:.4f}',
            r['severity'],
            f'{int(r["rule_violations"])}/10',
            f'{int(r["anomaly_votes"])}/3',
            r['top_feature_1_label'] if pd.notna(r['top_feature_1_label']) else '-',
        ))
    add_table(doc,
        ['Username', 'Score', 'Severity', 'Rules', 'Votes', 'Top SHAP Cause'],
        top_rows, 2, 'Top 5 anomalous users with SHAP explanations')

    add_heading_lvl2(doc, '5.3', 'Ablation study')
    add_text(doc,
        'To see what each member brings on its own, we run IF, LOF, and EE separately. Table '
        '3 collects their detection counts, how much they agree with one another, and their '
        'proxy-based quality scores.')

    add_table(doc,
        ['Model', 'Detected', 'Precision@K', 'Avg rule_viol', 'Avg score'],
        [(name, n, f'{pk:.4f}', f'{arv:.2f}', f'{asc:.4f}')
         for (name, n, pk, arv, asc) in IND_ABLATION],
        3, 'Ablation study: individual model performance')

    add_text(doc,
        'The three detectors clearly specialize. EE leads on Precision@K '
        f'({_ee[2]:.1%}) and on average rule violations ({_ee[3]:.2f}), staying close to the '
        'rule engine. LOF goes the opposite way '
        f'(avg rule_viol = {_lof[3]:.2f}, below the population baseline of {_BASE_RV:.2f}), a '
        'sign that it flags anomalies the rules say nothing about. Precisely that contrast is '
        'why a heterogeneous ensemble makes sense: no two models cover the same ground.')
    add_text(doc,
        'The agreement scores in Table 4 (Jaccard and Cohen’s Kappa) point the same way. LOF '
        f'and EE barely overlap (Kappa = {_lofee[2]:.2f}), so the user sets they flag are '
        'largely disjoint.')

    add_table(doc,
        ['Model Pair', 'Jaccard', 'Cohen’s Kappa', 'Interpretation'],
        [(pair, f'{j:.3f}', f'{k:.3f}', interp) for (pair, j, k, interp) in AGREEMENT],
        4, 'Inter-model agreement metrics')

    add_heading_lvl2(doc, '5.4', 'Ensemble configuration ablation')
    add_text(doc,
        'We then look past single models to every non-empty combination of the three — seven '
        'in all (three singletons, three pairs, and the full triple) — and ask whether one of '
        'them reliably beats the rest. To keep things fair, we fuse model scores by rank '
        '(which shrugs off scale differences between detectors) and let each configuration '
        'flag its top-K users, with K set to the number of proxy positives (rule_violations '
        '≥ 6). Accuracy, Precision, Recall, and F1 are then measured against the rule-based '
        'proxy, and Table 5 lays out the outcome.')

    combo_rows = []
    for name, acc, prec, rec, f1 in COMBO_ABLATION:
        combo_rows.append((name, f'{acc:.4f}', f'{prec:.4f}', f'{rec:.4f}', f'{f1:.4f}'))
    add_table(doc,
        ['Configuration', 'Accuracy', 'Precision', 'Recall', 'F1'],
        combo_rows, 5, 'Accuracy of seven ensemble configurations (rule-based proxy)')

    _sens_desc = ', '.join(f'{SENS_BEST[t][0]} (≥{t})' for t in _SENS_T)
    _ens_f1 = ', '.join(f'{SENS_TABLE["IF+LOF+EE"][t]:.3f}' for t in _SENS_T)
    add_text(doc,
        'A sensitivity analysis evaluates each configuration against proxy thresholds '
        f'rule_violations ≥ {_SENS_T[0]} … {_SENS_T[-1]} (best per threshold: {_sens_desc}). '
        'EE attains the highest F1 in most settings; however, this advantage is partly an '
        'artifact of the proxy being derived from the same rule engine, so EE detections '
        'align with rule violations by construction. Individual models are also volatile '
        'across definitions — for example, LOF is competitive at lower thresholds but '
        'collapses at higher ones. '
        + ('Critically, the full IF+LOF+EE ensemble is never the worst configuration at any '
           f'threshold (F1 = {_ens_f1}), remaining consistently mid-ranked. '
           if _ENS_NEVER_WORST else
           'The full ensemble stays competitive across thresholds. ')
        + 'This demonstrates that the value of the heterogeneous ensemble lies in robustness '
        'and stable generalization across differing anomaly definitions rather than in '
        'maximizing any single (rule-biased) proxy metric.')

    add_heading_lvl2(doc, '5.5', 'SHAP feature importance')
    add_text(doc,
        'Ranking the eleven features by mean absolute SHAP value (Fig. 2) puts rule '
        'violations clearly in front, with host diversity, shared device risk, and critical '
        'server access ratio close behind; each of these four sits above 0.47. Three features '
        'contribute nothing on this dataset (sensitive group membership, IP network risk, and '
        'privilege level all score zero), which likely reflects the make-up of this '
        'particular data rather than the features themselves and is worth re-checking on '
        'other deployments.')
    add_figure(doc, 'docs/diagrams/shap_importance.png', 2,
        'Mean absolute SHAP value for each feature, computed on the Isolation Forest '
        'sub-model and ordered from most to least important. Three of the features contribute '
        'nothing on this dataset.',
        width_cm=8.5, span_columns=False)
    add_text(doc,
        'Per-user SHAP top causes provide auditable explanations. For instance, the highest-'
        f'scoring user is flagged due to "{TOP5.iloc[0]["top_feature_1_label"]}", '
        'enabling security analysts to focus investigation efforts on specific behavioral '
        'patterns rather than generic anomaly labels.')

    add_heading_lvl2(doc, '5.6', 'Effect of class balancing (balanced vs. imbalanced)')
    add_text(doc,
        'Because the proxy positive class is rare (57 of 887 users, a 1:14 imbalance), we '
        'conducted a supplementary experiment to quantify how class balancing affects '
        'detectability. Models are trained on the original imbalanced split and on a '
        'generatively balanced split (minority oversampling with per-feature Gaussian jitter to '
        'a 50/50 ratio), and both are evaluated on the same realistic (imbalanced) test set.')
    add_text(doc,
        'A supervised probe makes the effect explicit. A linear classifier (logistic '
        'regression) collapses under imbalance (F1 = 0.000 — it labels every user normal to '
        'maximize accuracy) and only begins detecting anomalies after balancing (F1 = 0.286). '
        'A strong non-linear classifier (random forest) already reaches F1 = 0.970 without '
        'balancing and does not benefit from it (F1 = 0.919). Balancing therefore helps weak '
        'models but is unnecessary for strong ones; the very high random-forest score also '
        'signals circularity, since the rule-based proxy is largely a function of the same '
        'graph features.')
    add_text(doc,
        'When we re-train the three unsupervised detectors on balanced data and rerun the '
        'rank-fusion ablation, the methodological point becomes clear (Table 6). F1 jumps '
        'after balancing: the full IF+LOF+EE ensemble moves from 0.456 to 0.778 and EE from '
        '0.474 to 0.816, even as Accuracy falls from 0.930 to 0.778. The ~0.93 accuracy under '
        'imbalance is thus an illusion, produced by always predicting the majority "normal" '
        'class, and that is exactly why we treat F1 rather than accuracy as the headline '
        'number. On the balanced set the two coincide by construction: a 50/50 prior with a '
        'top-50% decision forces Precision = Recall = F1 = Accuracy.')
    add_table(doc,
        ['Configuration', 'Acc (imbal.)', 'F1 (imbal.)', 'Acc (bal.)', 'F1 (bal.)'],
        [
            ('IF',        '0.912', '0.316', '0.707', '0.707'),
            ('LOF',       '0.894', '0.175', '0.659', '0.659'),
            ('EE',        '0.932', '0.474', '0.816', '0.816'),
            ('IF+LOF',    '0.899', '0.211', '0.699', '0.699'),
            ('IF+EE',     '0.930', '0.456', '0.714', '0.714'),
            ('LOF+EE',    '0.930', '0.456', '0.784', '0.784'),
            ('IF+LOF+EE', '0.930', '0.456', '0.778', '0.778'),
        ], 6, 'Ensemble configurations under imbalanced vs. balanced training (rule-based proxy)')
    add_text(doc,
        'To rule out a lucky seed, we repeat the experiment across ten seeds (Table 7). In '
        'both settings the per-configuration standard deviation hovers around 0.01, well '
        'under the gaps that separate configurations. Because every detector is refit on each '
        'seed here, the imbalanced means come out a little different from the canonical '
        'stored-score values in Table 6 (EE is especially sensitive to refitting); the '
        'ranking itself, though, holds. We stress that the deployed pipeline stays '
        'unsupervised, and that balancing is only a supervised lens applied for this analysis. '
        'Its role is to justify reporting F1 over accuracy and to expose how imbalance '
        'inflates accuracy, not to change the pipeline we actually run.')
    add_table(doc,
        ['Configuration', 'F1 imbalanced (mean ± std)', 'F1 balanced (mean ± std)'],
        [
            ('IF',        '0.282 ± 0.015', '0.692 ± 0.019'),
            ('LOF',       '0.175 ± 0.000', '0.666 ± 0.008'),
            ('EE',        '0.368 ± 0.000', '0.806 ± 0.003'),
            ('IF+LOF',    '0.212 ± 0.006', '0.724 ± 0.025'),
            ('IF+EE',     '0.389 ± 0.014', '0.700 ± 0.006'),
            ('LOF+EE',    '0.333 ± 0.000', '0.798 ± 0.009'),
            ('IF+LOF+EE', '0.407 ± 0.020', '0.776 ± 0.010'),
        ], 7, '10-seed stability of F1 (detectors re-fitted per seed): imbalanced vs. balanced')

    add_heading_lvl2(doc, '5.7', 'Confusion matrix')
    add_text(doc,
        'Table 8 and Fig. 3 report the confusion matrix of the deployed ensemble flag — a user is flagged '
        'when at least two of three models agree or the fused score exceeds 0.75 — against the '
        f'rule-based proxy (rule_violations ≥ 6). Of {TOTAL_USERS:,} users, the system '
        f'flags {_CM_TP + _CM_FP} as anomalous; {_CM_TP} of these coincide with proxy positives '
        f'(true positives), with {_CM_FP} false positives, {_CM_FN} false negatives, and '
        f'{_CM_TN:,} true negatives — giving precision {_CM_P:.2f}, recall {_CM_R:.2f}, '
        f'F1 {_CM_F1:.2f}, and accuracy {_CM_ACC:.2f}. These values reflect the fixed operating '
        'point of the deployed system; the configuration ablation (Section 5.4) instead ranks '
        'the top-K users with K equal to the number of proxy positives, which raises the '
        'full-ensemble F1 to 0.46 by construction. Consistent with the rest of the '
        'evaluation, agreement is measured against a weak rule-based proxy rather than '
        'expert-validated labels; off-diagonal cells therefore include both genuine errors and '
        'unsupervised detections that are orthogonal to the encoded rules.')
    add_table(doc,
        ['', 'Proxy positive', 'Proxy negative'],
        [
            ('Flagged (predicted positive)', str(_CM_TP), str(_CM_FP)),
            ('Not flagged (predicted negative)', str(_CM_FN), str(_CM_TN)),
        ], 8, 'Confusion matrix: deployed ensemble flag vs. rule-based proxy')

    add_figure(doc, 'docs/diagrams/confusion_matrix.png', 3,
        'Heatmap of the confusion matrix comparing the deployed ensemble flag with the '
        'rule-based proxy, showing the TP, FP, FN, and TN counts.',
        width_cm=8.0, span_columns=False)

    add_heading_lvl2(doc, '5.8', 'Discussion')
    add_text(doc,
        'Three problems in AD anomaly detection shaped this design. The knowledge-graph view '
        'answers relational questions that a flat table simply cannot. The heterogeneous '
        'ensemble, by mixing inductive biases, catches more than one flavor of anomaly. And '
        'SHAP attributions together with quantile thresholds keep the output transparent and '
        'reproducible enough to drop into a security audit.')
    add_text(doc,
        'The obvious weakness is that we never had real ground-truth labels for an absolute '
        'precision/recall read. Our proxy (rule_violations ≥ 6) tilts toward detectors that '
        'echo the domain rules (EE) and is unkind to those that flag orthogonal anomalies '
        '(LOF), even when both are useful in practice. Closing that gap will mean bringing in '
        'expert-validated labels and studying how sensitive the results are to the threshold.')

    # ── 6. CONCLUSION ────────────────────────────────────────────────
    add_heading_lvl1(doc, 6, 'Conclusion')
    add_text(doc,
        'We have described a knowledge-graph approach to anomaly detection over Active '
        'Directory audit logs, one that ties together Neo4j, a rule-based knowledge engine, a '
        'heterogeneous ML ensemble (IF + LOF + EE), SHAP explainability, and a '
        'knowledge-base-grounded explainer (Phase 7) that writes up each anomaly in plain '
        'language with its sources attached (MITRE ATT&CK techniques and Windows Event IDs). '
        f'Run over {TOTAL_USERS:,} users from 1.8 million AD events, the pipeline marked '
        f'{TOTAL_ANOMALI} as anomalous at MEDIUM severity or above, spread across the '
        'CRITICAL, HIGH, and MEDIUM tiers, with quantile-based thresholds standing in for '
        'hand-picked cut-offs.')
    add_text(doc,
        'Our ablation showed the three components pulling in complementary directions: EE '
        'leans on the domain rules, LOF surfaces orthogonal local anomalies, and IF lands '
        f'between them. The low LOF-EE agreement (Cohen’s Kappa = {_lofee[2]:.2f}) is '
        'empirical support for combining them. A class-balancing analysis then exposed the '
        'imbalanced accuracy (~0.93) as misleading and established F1 — which climbs from '
        '0.46 to 0.78 once the classes are balanced — as the metric worth quoting. SHAP, '
        'finally, turns opaque anomaly scores into evidence, which is what makes the system '
        'usable day to day in security operations.')
    add_text(doc,
        'Future work will focus on (1) expert-validated ground truth for absolute '
        'precision/recall evaluation, (2) threshold sensitivity analysis with ROC and '
        'cost-sensitive metrics, and (3) extension to temporal anomaly detection using '
        'sequence models (e.g., LSTM autoencoders) integrated with the graph representation.')

    # ── CONFLICTS OF INTEREST ────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Conflicts of interest')
    set_font(run, size=11, bold=True)
    add_text(doc, 'The authors declare no conflict of interest.', indent=False)

    # ── AUTHOR CONTRIBUTIONS ─────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Author contributions')
    set_font(run, size=11, bold=True)
    add_text(doc,
        'Conceptualization, methodology, software, validation, formal analysis, '
        'investigation, data curation, writing—original draft preparation, '
        'writing—review and editing, visualization: Mahathir Muhammad.', indent=False)

    # ── REFERENCES ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('References')
    set_font(run, size=12, bold=True)

    references = [
        '[1] P.-F. Marteau, “Random partitioning forest for point-wise and collective anomaly detection — Application to network intrusion detection”, IEEE Transactions on Information Forensics and Security, Vol. 16, pp. 2157–2172, 2021, doi: 10.1109/TIFS.2021.3050605.',
        '[2] E. Tufan, C. Tezcan, and C. Acarturk, “Anomaly-based intrusion detection by machine learning: A case study on probing attacks to an institutional network”, IEEE Access, Vol. 9, pp. 50078–50092, 2021, doi: 10.1109/ACCESS.2021.3068961.',
        '[3] H. Xu, G. Pang, Y. Wang, and Y. Wang, “Deep Isolation Forest for anomaly detection”, IEEE Transactions on Knowledge and Data Engineering, Vol. 35, No. 12, pp. 12591–12604, 2023, doi: 10.1109/TKDE.2023.3270293.',
        '[4] M. Landauer, S. Onder, F. Skopik, and M. Wurzenberger, “Deep learning for anomaly detection in log data: A survey”, Machine Learning with Applications, Vol. 12, Art. 100470, 2023, doi: 10.1016/j.mlwa.2023.100470.',
        '[5] V. H. Le and H. Zhang, “Log-based anomaly detection without log parsing”, In: Proc. of 36th IEEE/ACM Int. Conf. on Automated Software Engineering (ASE), pp. 492–504, 2021, doi: 10.1109/ASE51524.2021.9678773.',
        '[6] C. Smiliotopoulos, G. Kambourakis, and K. Barbatsalou, “On the detection of lateral movement through supervised machine learning and an open-source tool to create turnkey datasets from Sysmon logs”, International Journal of Information Security, Vol. 22, No. 6, pp. 1893–1919, 2023, doi: 10.1007/s10207-023-00725-8.',
        '[7] X. Sun and J. Yang, “HetGLM: Lateral movement detection by discovering anomalous links with heterogeneous graph neural network”, In: Proc. of IEEE Int. Performance, Computing, and Communications Conf. (IPCCC), pp. 404–411, 2022, doi: 10.1109/IPCCC55026.2022.9894347.',
        '[8] Z. Jia, Y. Xiong, Y. Nan, Y. Zhang, J. Zhao, and M. Wen, “MAGIC: Detecting advanced persistent threats via masked graph representation learning”, In: Proc. of 33rd USENIX Security Symposium, Philadelphia, PA, USA, 2024. arXiv:2310.09831.',
        '[9] A. Zhou, X. Xu, R. Raghunathan, A. Lal, X. Guan, B. Yu, and B. Li, “KnowGraph: Knowledge-enabled anomaly detection via logical reasoning on graph data”, In: Proc. of ACM SIGSAC Conf. on Computer and Communications Security (CCS), 2024, doi: 10.1145/3658644.3690354.',
        '[10] F. Xiao, S. Chen, S. Chen, Y. Ma, H. He, and J. Yang, “SENTINEL: Insider threat detection based on multi-timescale user behavior interaction graph learning”, IEEE Transactions on Network Science and Engineering, Vol. 12, No. 2, pp. 774–790, 2025, doi: 10.1109/TNSE.2024.3519155.',
        '[11] K. C. Roy and G. Chen, “GraphCH: A deep framework for assessing cyber-human aspects in insider threat detection”, IEEE Transactions on Dependable and Secure Computing, Vol. 21, No. 5, pp. 4495–4509, 2024, doi: 10.1109/TDSC.2024.3353929.',
        '[12] F. R. Alzaabi and A. Mehmood, “A review of recent advances, challenges, and opportunities in malicious insider threat detection using machine learning methods”, IEEE Access, Vol. 12, pp. 30907–30927, 2024, doi: 10.1109/ACCESS.2024.3369906.',
        '[13] D. Gaspar, P. Silva, and C. Silva, “Explainable AI for intrusion detection systems: LIME and SHAP applicability on multi-layer perceptron”, IEEE Access, Vol. 12, pp. 30164–30175, 2024, doi: 10.1109/ACCESS.2024.3368377.',
        '[14] P. Barnard, N. Marchetti, and L. A. DaSilva, “Robust network intrusion detection through explainable artificial intelligence (XAI)”, IEEE Networking Letters, Vol. 4, No. 3, pp. 167–171, 2022, doi: 10.1109/LNET.2022.3186589.',
        '[15] A. Nascita, G. Aceto, D. Ciuonzo, A. Montieri, V. Persico, and A. Pescapè, “A survey on explainable artificial intelligence for Internet traffic classification and prediction, and intrusion detection”, IEEE Communications Surveys & Tutorials, Vol. 27, No. 5, pp. 3165–3198, 2025, doi: 10.1109/COMST.2024.3504955.',
        '[16] I. M. Sayem, M. I. Sayed, S. Saha, and A. Haque, “ENIDS: A deep learning-based ensemble framework for network intrusion detection systems”, IEEE Transactions on Network and Service Management, Vol. 21, No. 5, pp. 5809–5825, 2024, doi: 10.1109/TNSM.2024.3414305.',
        '[17] M. Lin, K. Yang, Z. Yu, Y. Shi, and C. L. P. Chen, “Hybrid ensemble broad learning system for network intrusion detection”, IEEE Transactions on Industrial Informatics, Vol. 20, No. 4, pp. 5622–5633, 2024, doi: 10.1109/TII.2023.3332957.',
        '[18] O. Arreche, I. Bibers, and M. Abdallah, “A two-level ensemble learning framework for enhancing network intrusion detection systems”, IEEE Access, Vol. 12, pp. 83830–83857, 2024, doi: 10.1109/ACCESS.2024.3407029.',
        '[19] A. K. Mananayaka and S. S. Chung, “Network intrusion detection with two-phased hybrid ensemble learning and automatic feature selection”, IEEE Access, Vol. 11, pp. 45154–45167, 2023, doi: 10.1109/ACCESS.2023.3274474.',
        '[20] M. Vishwakarma and N. Kesswani, “StaEn-IDS: An explainable stacking ensemble deep neural network-based intrusion detection system for IoT”, IEEE Access, Vol. 13, pp. 109713–109728, 2025, doi: 10.1109/ACCESS.2025.3582391.',
        '[21] L. Mahmoud, M. Liyanage, J. Singla, and S. Gangopadhyay, “DSEM-NIDS: Enhanced network intrusion detection system using deep stacking ensemble model”, IEEE Open Journal of the Computer Society, Vol. 6, pp. 955–967, 2025, doi: 10.1109/OJCS.2025.3581036.',
        '[22] H. Zhang, K. Zeng, and S. Lin, “Federated graph neural network for fast anomaly detection in controller area networks”, IEEE Transactions on Information Forensics and Security, Vol. 18, pp. 1566–1579, 2023, doi: 10.1109/TIFS.2023.3240291.',
        '[23] G. Duan, H. Lv, H. Wang, and G. Feng, “Application of a dynamic line graph neural network for intrusion detection with semisupervised learning”, IEEE Transactions on Information Forensics and Security, Vol. 18, pp. 699–714, 2023, doi: 10.1109/TIFS.2022.3228493.',
        '[24] Z. Li, R. Liu, D. Chen, and Q. Hu, “OR-gate mixup multiscale spectral graph neural network for node anomaly detection”, IEEE Transactions on Neural Networks and Learning Systems, Vol. 36, No. 9, pp. 16692–16705, 2025, doi: 10.1109/TNNLS.2025.3569413.',
        '[25] S. Rekik and S. Mehmood, “A hybrid graph neural network and neural ODE model to intrusion detection in dynamic network topologies”, IEEE Access, Vol. 13, pp. 198201–198227, 2025, doi: 10.1109/ACCESS.2025.3635385.',
        '[26] J. Xiao, L. Yang, F. Zhong, X. Wang, H. Chen, and D. Li, “Robust anomaly-based insider threat detection using graph neural network”, IEEE Transactions on Network and Service Management, Vol. 20, No. 3, pp. 3717–3733, 2023, doi: 10.1109/TNSM.2022.3222635.',
        '[27] A. Yayla, L. Haghnegahdar, and E. Dincelli, “Explainable artificial intelligence for smart grid intrusion detection systems”, IT Professional, Vol. 24, No. 5, pp. 18–24, 2022, doi: 10.1109/MITP.2022.3163731.',
        '[28] A. Khan, M. A. Hussain, and F. Anwer, “A hybrid lightweight deep learning-based intrusion detection approach in IoT utilizing feature selection and explainable artificial intelligence”, IEEE Access, Vol. 13, pp. 192451–192466, 2025, doi: 10.1109/ACCESS.2025.3619449.',
        '[29] S. F. Akintade, K. Roy, and S.-T. Kim, “Hybrid deep machine learning feature selection for high-dimensional cybersecurity data”, IEEE Access, Vol. 13, pp. 172136–172156, 2025, doi: 10.1109/ACCESS.2025.3615582.',
        '[30] E.-U.-H. Qazi, T. Zia, M. H. Faheem, K. Shahzad, M. Imran, and Z. Ahmed, “Zero-touch network security (ZTNS): A network intrusion detection system based on deep learning”, IEEE Access, Vol. 12, pp. 141625–141638, 2024, doi: 10.1109/ACCESS.2024.3466470.',
        '[31] J. Saikam and K. Ch, “EESNN: Hybrid deep learning empowered spatial–temporal features for network intrusion detection system”, IEEE Access, Vol. 12, pp. 15930–15945, 2024, doi: 10.1109/ACCESS.2024.3350197.',
    ]
    _doi_re = re.compile(r'(doi:\s*)(10\.\S+?)\.?$')
    _arxiv_re = re.compile(r'(arXiv:)(\S+?)\.?$')
    _refnum_re = re.compile(r'^\[(\d+)\]')
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _rn = _refnum_re.match(ref)
        if _rn:
            add_bookmark_to_paragraph(p, f'ref_{_rn.group(1)}', int(_rn.group(1)))
        _m = _doi_re.search(ref)
        _a = _arxiv_re.search(ref)
        if _m:                       # link the DOI to https://doi.org/<doi>
            run = p.add_run(ref[:_m.start()] + _m.group(1))
            set_font(run, size=11)
            add_hyperlink(p, f'https://doi.org/{_m.group(2)}', _m.group(2), size=11)
            tail = p.add_run('.')
            set_font(tail, size=11)
        elif _a:                     # link the arXiv id to https://arxiv.org/abs/<id>
            run = p.add_run(ref[:_a.start()] + _a.group(1))
            set_font(run, size=11)
            add_hyperlink(p, f'https://arxiv.org/abs/{_a.group(2)}', _a.group(2), size=11)
            tail = p.add_run('.')
            set_font(tail, size=11)
        else:
            run = p.add_run(ref)
            set_font(run, size=11)

    out_path = 'output/IJIES_Draft_Paper_Mahathir_Muhammad.docx'
    doc.save(out_path)
    print(f'[OK] Paper saved: {out_path}')
    return out_path


if __name__ == '__main__':
    print('Generating IJIES draft paper...')
    path = build_paper()
    print(f'Done: {path}')
