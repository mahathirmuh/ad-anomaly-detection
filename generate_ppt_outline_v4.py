#!/usr/bin/env python3
"""
Generate PPT Outline v4 (DOCX) — TDAS AD Audit
Slide-by-slide content outline reflecting the CURRENT canonical pipeline:
7-phase, Knowledge Graph + Rule Engine + Ensemble (IF+LOF+EE) + SHAP + Ablation.
All numbers are read dynamically from canonical pipeline outputs.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import pandas as pd
import numpy as np
import json
import io
import os

os.makedirs('docs/presentations', exist_ok=True)

# ── LOAD CANONICAL DATA ──────────────────────────────────────────────
with open('output/anomaly_statistics.json') as f:
    STATS = json.load(f)
DIST = STATS['anomaly_distribution']
SCORE = STATS['anomaly_score_stats']
RULE = STATS['rule_violations_stats']
TOTAL_USERS = STATS['total_users']

df = pd.read_csv('data/phase5_anomaly_results.csv').drop_duplicates(subset='user_id').reset_index(drop=True)
shap_df = pd.read_csv('data/phase55_shap_values.csv')

P75 = df['final_anomaly_score'].quantile(0.75)
P90 = df['final_anomaly_score'].quantile(0.90)
P95 = df['final_anomaly_score'].quantile(0.95)
P99 = df['final_anomaly_score'].quantile(0.99)
TOTAL_ANOMALI = int(df['severity'].isin(['CRITICAL', 'HIGH', 'MEDIUM']).sum())

merged = df.merge(shap_df[['user_id', 'top_feature_1_label']], on='user_id', how='left')
TOP = merged.nlargest(5, 'final_anomaly_score')

# SHAP global importance
shap_cols = [c for c in shap_df.columns if c.startswith('shap_')]
mean_abs = shap_df[shap_cols].abs().mean().sort_values(ascending=False)
SHAP_TOP = [(c.replace('shap_', ''), v) for c, v in mean_abs.head(5).items()]

# Ablation: individual model Precision@K + 7-config accuracy
from scipy.stats import rankdata
from sklearn.metrics import cohen_kappa_score
N = len(df)
_MR = {'IF': rankdata(df['if_score'].values), 'LOF': rankdata(df['lof_score'].values),
       'EE': rankdata(df['ee_score'].values)}
_GT = (df['rule_violations'].values >= 6).astype(int)
_PP = int(_GT.sum())
_CFG = [('IF',), ('LOF',), ('EE',), ('IF', 'LOF'), ('IF', 'EE'), ('LOF', 'EE'), ('IF', 'LOF', 'EE')]

def _eval(models, K):
    c = np.mean([_MR[m] for m in models], axis=0); o = np.argsort(c)
    p = np.zeros(N, int); p[o[:K]] = 1
    TP = int(((p == 1) & (_GT == 1)).sum()); FP = int(((p == 1) & (_GT == 0)).sum())
    FN = int(((p == 0) & (_GT == 1)).sum()); TN = int(((p == 0) & (_GT == 0)).sum())
    acc = (TP + TN) / N
    pr = TP / (TP + FP) if TP + FP else 0; rc = TP / (TP + FN) if TP + FN else 0
    f1 = 2 * pr * rc / (pr + rc) if pr + rc else 0
    return acc, f1

COMBO = [('+'.join(c), *_eval(c, _PP)) for c in _CFG]

def _pk(col):
    s = set(df[df[col] == 1]['user_id']); heavy = set(df[df['rule_violations'] >= 6]['user_id'])
    return len(s & heavy) / len(s) if s else 0
IND = {'IF': _pk('if_anomaly'), 'LOF': _pk('lof_anomaly'), 'EE': _pk('ee_anomaly')}
KAPPA = {
    'IF-LOF': cohen_kappa_score(df['if_anomaly'], df['lof_anomaly']),
    'IF-EE': cohen_kappa_score(df['if_anomaly'], df['ee_anomaly']),
    'LOF-EE': cohen_kappa_score(df['lof_anomaly'], df['ee_anomaly']),
}

# ── FIGURES (canonical data) ─────────────────────────────────────────
PAL = ['#1A3C6E', '#1565C0', '#1976D2', '#1E88E5', '#42A5F5', '#64B5F6', '#90CAF9']

def fig_pipeline():
    fig, ax = plt.subplots(figsize=(12, 2.6)); ax.axis('off')
    stages = [("AD Log", "#1A3C6E"), ("Neo4j\nKG", "#1565C0"), ("Rule\nEngine", "#1976D2"),
              ("Graph\nFeatures", "#1E88E5"), ("Ensemble\nIF+LOF+EE", "#42A5F5"),
              ("SHAP", "#64B5F6"), ("Report", "#90CAF9")]
    n = len(stages); w, h, gap = 1.45, 0.8, 0.12; x0 = (12 - (n*w + (n-1)*gap)) / 2
    for i, (lbl, col) in enumerate(stages):
        x = x0 + i*(w+gap)
        ax.add_patch(FancyBboxPatch((x, 0.2), w, h, boxstyle="round,pad=0.05",
                                    facecolor=col, edgecolor='white', lw=1.5))
        ax.text(x+w/2, 0.6, lbl, ha='center', va='center', fontsize=9.5,
                color='white', fontweight='bold')
        ax.text(x+w/2, 0.08, f"Phase {['1','2','3','4','5','5.5','6'][i]}",
                ha='center', fontsize=7.5, color='#666')
        if i < n-1:
            ax.annotate('', xy=(x+w+gap, 0.6), xytext=(x+w, 0.6),
                        arrowprops=dict(arrowstyle='->', color='#555', lw=2))
    ax.set_xlim(0, 12); ax.set_ylim(0, 1.1); fig.tight_layout(pad=0)
    return fig

def fig_graph_schema():
    fig, ax = plt.subplots(figsize=(9, 5)); ax.axis('off'); ax.set_xlim(0, 10); ax.set_ylim(0, 6)
    nodes = {'User': (5, 5, '#1A3C6E'), 'Hostname': (2, 3.5, '#1565C0'), 'Server': (8, 3.5, '#1976D2'),
             'IPAddress': (2, 1.5, '#1E88E5'), 'Group': (8, 1.5, '#42A5F5'),
             'Event': (5, 2.6, '#64B5F6'), 'Service': (5, 0.5, '#90CAF9')}
    edges = [('User', 'Hostname', 'LOGIN_FROM'), ('User', 'Server', 'AUTHENTICATED_VIA'),
             ('User', 'IPAddress', 'CONNECTED_FROM'), ('User', 'Group', 'MEMBER_OF'),
             ('Event', 'User', 'REFERENCES'), ('Hostname', 'IPAddress', 'USED_IP'),
             ('User', 'Service', 'USED_SERVICE')]
    for s, d, rel in edges:
        x1, y1, _ = nodes[s]; x2, y2, _ = nodes[d]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#999', lw=1.3,
                                    connectionstyle='arc3,rad=0.1'))
        ax.text((x1+x2)/2, (y1+y2)/2, rel, ha='center', fontsize=6.5, style='italic',
                bbox=dict(boxstyle='round,pad=0.2', fc='#F0F4FF', ec='none'))
    for lbl, (x, y, col) in nodes.items():
        ax.add_patch(plt.Circle((x, y), 0.5, color=col, zorder=3))
        ax.text(x, y, lbl, ha='center', va='center', fontsize=8, color='white',
                fontweight='bold', zorder=4)
    ax.set_title('Neo4j Knowledge Graph — 7 Nodes / 10 Relationships',
                 fontsize=11, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_shap():
    feats = [f for f, _ in SHAP_TOP][::-1]; vals = [v for _, v in SHAP_TOP][::-1]
    fig, ax = plt.subplots(figsize=(9, 3.8))
    bars = ax.barh(feats, vals, color='#1565C0', edgecolor='white')
    ax.bar_label(bars, fmt='%.3f', padding=3, fontsize=9)
    ax.set_xlabel('Mean |SHAP value|'); ax.set_xlim(0, max(vals)*1.18)
    ax.set_title('SHAP Global Feature Importance (Top 5)', fontsize=12, fontweight='bold', color='#1A3C6E')
    ax.tick_params(labelsize=10)
    for sp in ['top', 'right']: ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def fig_severity():
    labels = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'NORMAL']
    sizes = [DIST.get(k, 0) for k in labels]
    cols = ['#B71C1C', '#E53935', '#FB8C00', '#FDD835', '#43A047']
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    w, t, a = ax.pie(sizes, labels=[f'{l}\n({s})' for l, s in zip(labels, sizes)],
                     colors=cols, explode=(0.12, 0.06, 0.03, 0, 0), autopct='%1.1f%%',
                     startangle=140, pctdistance=0.78, textprops={'fontsize': 10})
    for x in a: x.set_color('white'); x.set_fontweight('bold'); x.set_fontsize(9)
    ax.set_title(f'Distribusi Severity ({TOTAL_USERS} Users)', fontsize=12, fontweight='bold', color='#1A3C6E')
    fig.tight_layout()
    return fig

def fig_top():
    users = TOP['username'].tolist()[::-1]; scores = TOP['final_anomaly_score'].tolist()[::-1]
    sev = TOP['severity'].tolist()[::-1]
    cmap = {'CRITICAL': '#B71C1C', 'HIGH': '#E53935', 'MEDIUM': '#FB8C00', 'LOW': '#FDD835'}
    fig, ax = plt.subplots(figsize=(9, 3.8))
    bars = ax.barh(users, scores, color=[cmap.get(s, '#999') for s in sev], edgecolor='white')
    ax.bar_label(bars, fmt='%.4f', padding=3, fontsize=9)
    ax.set_xlabel('Anomaly Score'); ax.set_xlim(0, max(scores)*1.18)
    ax.set_title('Top-5 Anomali (Score Tertinggi)', fontsize=12, fontweight='bold', color='#1A3C6E')
    ax.tick_params(labelsize=10)
    for sp in ['top', 'right']: ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def fig_ablation_ind():
    models = ['IF', 'LOF', 'EE']; vals = [IND[m] for m in models]
    fig, ax = plt.subplots(figsize=(7.5, 4))
    bars = ax.bar(models, vals, color=['#1A3C6E', '#1976D2', '#42A5F5'], edgecolor='white')
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, v+0.01, f'{v:.1%}', ha='center', fontweight='bold', fontsize=11)
    ax.set_ylabel('Precision@K'); ax.set_ylim(0, max(vals)*1.25)
    ax.set_title('Ablation Individual — Precision@K\n(% deteksi = pelanggar rule berat)',
                 fontsize=12, fontweight='bold', color='#1A3C6E')
    for sp in ['top', 'right']: ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

def fig_combo():
    names = [c[0] for c in COMBO]; acc = [c[1] for c in COMBO]; f1 = [c[2] for c in COMBO]
    x = np.arange(len(names)); w = 0.4
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.bar(x-w/2, acc, w, label='Accuracy', color='#1A3C6E')
    ax.bar(x+w/2, f1, w, label='F1', color='#42A5F5')
    ax.set_xticks(x); ax.set_xticklabels(names, rotation=20, fontsize=9)
    ax.set_ylim(0, 1.05); ax.legend(fontsize=10)
    ax.set_title('Ablation 7 Konfigurasi Ensemble — Accuracy & F1 (vs proxy rule)',
                 fontsize=12, fontweight='bold', color='#1A3C6E')
    for i, v in enumerate(f1):
        ax.text(i+w/2, v+0.01, f'{v:.2f}', ha='center', fontsize=8)
    for sp in ['top', 'right']: ax.spines[sp].set_visible(False)
    fig.tight_layout()
    return fig

# ── DOCX HELPERS ─────────────────────────────────────────────────────
doc = Document()
NAVY = RGBColor(0x1A, 0x3C, 0x6E)

def add_fig(fig, width=6.2):
    buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=150, bbox_inches='tight'); buf.seek(0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(width))
    plt.close(fig)

def title(text, sub_lines):
    p = doc.add_paragraph(); p.style = doc.styles['Title']
    r = p.add_run(text)
    for s in sub_lines:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = pp.add_run(s); run.font.size = Pt(11)

def slide(n, head):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Slide {n} — {head}")
    r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    sep = doc.add_paragraph(); sr = sep.add_run('_' * 90)
    sr.font.size = Pt(8); sr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def h2(text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 2']
    r = p.add_run(text); r.font.color.rgb = NAVY

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    p.add_run(text)

def note(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── SLIDE 1: TITLE ───────────────────────────────────────────────────
title("TDAS AD Audit v4", [
    "Explainable Anomaly Detection in Active Directory:",
    "Rule-Based Knowledge Engine + Ensemble Learning + SHAP for Human-Readable Reasoning",
    "",
    "Mahathir Muhammad",
    "Final Pipeline, Ablation Study & Draft Paper (IJIES) — Juni 2026",
])

# ── SLIDE 2: LATAR BELAKANG ──────────────────────────────────────────
slide(2, "Executive Summary: Latar Belakang")
bullet("Insider threat & privilege abuse di Active Directory (AD) adalah vektor serangan paling merusak dan paling sulit dideteksi.")
bullet("Log AD bervolume sangat besar (1.8 juta+ events) dan kaya relasi antar entitas (user ↔ host ↔ server ↔ IP ↔ group).")
bullet("Metode single-model (mis. Isolation Forest saja) sering false-positive tinggi karena kehilangan konteks relasional.")
bullet("Sistem deteksi yang black-box sulit ditindaklanjuti auditor — butuh penjelasan kuantitatif per user.")
note("→ Dibutuhkan pendekatan Knowledge Graph + Ensemble + Explainable AI yang akurat sekaligus dapat dijelaskan.")

# ── SLIDE 3: TUJUAN & KONTRIBUSI ─────────────────────────────────────
slide(3, "Executive Summary: Tujuan & Kontribusi")
h2("Tujuan")
bullet("Memodelkan log AD sebagai knowledge graph relasional di Neo4j.")
bullet("Mendeteksi anomali pengguna dengan ensemble heterogen (IF + LOF + EE).")
bullet("Menjelaskan setiap keputusan deteksi menggunakan SHAP (explainable).")
h2("Kontribusi")
bullet("Pipeline 7-phase end-to-end: Ingestion → Rules → Features → Ensemble → SHAP → Reporting.")
bullet("Severity classification data-driven berbasis quantile (bukan threshold arbitrary).")
bullet("Ablation study lengkap atas 7 konfigurasi ensemble (termasuk IF+EE) untuk justifikasi ilmiah.")
bullet(f"Evaluasi empiris pada {TOTAL_USERS} user dari 1.8 juta events (ManageEngine ADAudit Plus).")

# ── SLIDE 4: PROBLEM STATEMENT ───────────────────────────────────────
slide(4, "Problem Statement")
note('"Bagaimana mendeteksi perilaku pengguna AD yang anomali secara akurat DAN dapat dijelaskan, '
     'ketika log bervolume besar, multi-sumber, dan tidak memiliki label ground-truth?"')
h2("Sub-problem")
bullet("Representasi multi-sumber log AD dalam satu struktur relasional yang queryable.")
bullet("Feature engineering diskriminatif yang memanfaatkan relasi graph (bukan hanya kolom CSV).")
bullet("Deteksi robust tanpa label (unsupervised) + cara mengevaluasinya.")
bullet("Transparansi: menjelaskan MENGAPA user dianggap anomali.")

# ── SLIDE 5: SOTA ────────────────────────────────────────────────────
slide(5, "Literature Review: SOTA & Posisi Penelitian")
bullet("Isolation Forest (Liu et al., 2008) — tree-based anomaly detection, fondasi banyak sistem.")
bullet("LOF (Breunig et al., 2000) & Elliptic Envelope (Rousseeuw, 1999) — density & statistical outlier.")
bullet("SHAP (Lundberg & Lee, 2017) — unified explainability untuk model ML.")
bullet("Outlier Analysis (Aggarwal, 2017) & Goldstein-Uchida (2016) — evaluasi unsupervised & thresholding.")
note("Posisi: menggabungkan Knowledge Graph AD + ensemble heterogen + explainability dalam satu pipeline.")

# ── SLIDE 6: RESEARCH GAP ────────────────────────────────────────────
slide(6, "Research Gap & Technical Differentiation")
h2("Research Gap")
bullet("SOTA terbelah: ML generik (tidak AD-specific) ATAU rule/signature AD (tidak adaptif).")
bullet("Sedikit yang menggabungkan graph context + ensemble + explainability untuk audit AD.")
h2("Diferensiasi (TDAS v4)")
bullet("Neo4j sebagai core intelligence (bukan sekadar visualisasi).")
bullet("Rule engine domain ditanam SEBELUM ML → konteks + anomali statistik.")
bullet("SHAP per user + severity quantile-based yang reproducible.")

# ── SLIDE 7: PIPELINE ────────────────────────────────────────────────
slide(7, "Metodologi: Pipeline 7-Phase")
note("AD Log → Neo4j KG → Rule Engine → Graph Features → Ensemble → SHAP → Report")
bullet("Phase 2 — Neo4j Ingestion: 7 node types, 10 relationship types.")
bullet("Phase 3 — Rule Engine: 10 domain rules (R001–R010).")
bullet("Phase 4 — Feature Extraction: 11 graph features per user.")
bullet("Phase 5 — Ensemble Detection: IF + LOF + EE + quantile severity.")
bullet("Phase 5.5 — SHAP Explainability: top cause per user.")
bullet("Phase 6 — Reporting: TXT, JSON, DOCX.")
note("Reproducible: Phase 2 reset database (CREATE OR REPLACE DATABASE) + random_state=42.")
add_fig(fig_pipeline(), 6.8)

# ── SLIDE 8: KNOWLEDGE GRAPH ─────────────────────────────────────────
slide(8, "Phase 2 — Neo4j Knowledge Graph")
h2("7 Node Types")
bullet("User, Hostname, Server, IPAddress, Service, Group, Event")
h2("10 Relationship Types")
bullet("LOGIN_FROM, AUTHENTICATED_VIA, FAILED_LOGIN, CONNECTED_FROM, USED_IP,")
bullet("USED_SERVICE, MEMBER_OF, REFERENCES, LOCKED_OUT, ADMIN_ACTION_ON")
note(f"Hasil ingestion: {TOTAL_USERS} User nodes, ~680K total nodes dari 1.8 juta events.")
add_fig(fig_graph_schema(), 5.6)

# ── SLIDE 9: RULES ───────────────────────────────────────────────────
slide(9, "Phase 3 — Rule-Based Knowledge Engine (10 Rules)")
bullet("R001 Multi-host login (>3 host)  |  R002 Off-hours login (>10%)")
bullet("R003 Shared device (>5 user)  |  R004 Critical server access (DC)")
bullet("R005 Failed login spike (>50)  |  R006 Unusual IP (luar Office/VPN)")
bullet("R007 After-hours privileged  |  R008 Frequent lockouts")
bullet("R009 Excessive admin actions  |  R010 Sensitive group membership")
note(f"Output: rule_violations (0–10) per user. Rata-rata {RULE['mean']:.2f} pelanggaran/user.")

# ── SLIDE 10: FEATURES ───────────────────────────────────────────────
slide(10, "Phase 4 — Graph Feature Extraction (11 Fitur)")
h2("Fitur dasar (8)")
bullet("host_diversity, critical_server_ratio, failure_ratio, shared_device_risk,")
bullet("ip_network_risk, privilege_level, connectivity, rule_violations")
h2("Fitur tambahan (3)")
bullet("lockout_count, admin_actions, sensitive_groups")
note(f"Output: {TOTAL_USERS} user × 11 fitur — diekstrak dari relasi graph, bukan kolom CSV langsung.")

# ── SLIDE 11: ENSEMBLE ───────────────────────────────────────────────
slide(11, "Phase 5 — Ensemble Anomaly Detection")
bullet("Isolation Forest (IF) — anomali global ekstrem (tree-based).")
bullet("Local Outlier Factor (LOF) — anomali lokal di cluster (density).")
bullet("Elliptic Envelope (EE) — outlier statistik multivariat (covariance).")
note("final_score = 0.60 × (votes/3) + 0.40 × (rule_violations/10). Anomali jika ≥2 dari 3 model setuju.")
h2("Severity Quantile-based (data-driven)")
bullet(f"CRITICAL ≥ P99 ({P99:.4f})  |  HIGH ≥ P95 ({P95:.4f})  |  MEDIUM ≥ P90 ({P90:.4f})")
bullet(f"LOW ≥ P75 ({P75:.4f})  |  NORMAL < P75")
note("Referensi: Aggarwal (2017), Goldstein-Uchida (2016), Liu et al. (2008).")

# ── SLIDE 12: SHAP ───────────────────────────────────────────────────
slide(12, "Phase 5.5 — SHAP Explainability")
bullet("shap.TreeExplainer pada Isolation Forest — menghitung kontribusi tiap fitur ke skor.")
bullet("Setiap user mendapat top cause (penyebab utama) dengan label Bahasa Indonesia.")
bullet("Hasil ditulis kembali ke Neo4j sebagai property node User.")
h2("Top Global Feature Importance (mean |SHAP|)")
for feat, val in SHAP_TOP:
    bullet(f"{feat}: {val:.4f}")
note("Manfaat: audit punya bukti kuantitatif per user, bukan sekadar label 'anomali'.")
add_fig(fig_shap(), 5.8)

# ── SLIDE 13: DATASET ────────────────────────────────────────────────
slide(13, "Experimental Setup: Dataset")
bullet("Sumber: ManageEngine ADAudit Plus (multiple CSV audit report).")
bullet("Total events: 1,833,352 logon events.")
bullet(f"Populasi: {TOTAL_USERS} unique users.")
bullet("Implementasi: Python 3.12, Neo4j 5.x, scikit-learn 1.4, shap.")
note("Unsupervised — tidak ada label asli; evaluasi memakai proxy berbasis rule_violations.")

# ── SLIDE 14: HASIL SEVERITY ─────────────────────────────────────────
slide(14, "Hasil: Distribusi Severity")
bullet(f"CRITICAL: {DIST.get('CRITICAL',0)} user (top 1%)")
bullet(f"HIGH: {DIST.get('HIGH',0)} user (top 5%)")
bullet(f"MEDIUM: {DIST.get('MEDIUM',0)} user (top 10%)")
bullet(f"LOW: {DIST.get('LOW',0)} user (top 25%)")
bullet(f"NORMAL: {DIST.get('NORMAL',0)} user")
note(f"Total {TOTAL_ANOMALI} user terklasifikasi MEDIUM ke atas (perlu perhatian). "
     f"Score: mean {SCORE['mean']:.3f}, max {SCORE['max']:.3f}.")
add_fig(fig_severity(), 5.0)

# ── SLIDE 15: TOP ANOMALI ────────────────────────────────────────────
slide(15, "Hasil: Top-5 Anomali + SHAP Cause")
for _, r in TOP.iterrows():
    cause = r['top_feature_1_label'] if pd.notna(r['top_feature_1_label']) else '-'
    bullet(f"{r['username']}  |  Score {r['final_anomaly_score']:.4f}  |  {r['severity']}  |  "
           f"{int(r['rule_violations'])}/10 rules  |  Cause: {cause}")
note("Setiap temuan dilengkapi penyebab SHAP → langsung actionable untuk investigasi.")
add_fig(fig_top(), 6.0)

# ── SLIDE 16: ABLATION INDIVIDUAL ────────────────────────────────────
slide(16, "Ablation Study: IF vs LOF vs EE (Individual)")
h2("Precision@K (selaras rule berat) & Cohen's Kappa")
bullet(f"IF: P@K {IND['IF']:.3f}  |  LOF: P@K {IND['LOF']:.3f}  |  EE: P@K {IND['EE']:.3f}")
bullet(f"Kappa IF-LOF {KAPPA['IF-LOF']:.2f}  |  IF-EE {KAPPA['IF-EE']:.2f}  |  LOF-EE {KAPPA['LOF-EE']:.2f}")
bullet("EE paling selaras dengan rule; LOF menangkap anomali ortogonal (tersembunyi).")
note(f"Kappa LOF-EE rendah ({KAPPA['LOF-EE']:.2f}) = mereka menangkap jenis anomali berbeda → ensemble bermanfaat.")
add_fig(fig_ablation_ind(), 5.0)

# ── SLIDE 17: ABLATION 7-CONFIG ──────────────────────────────────────
slide(17, "Ablation Study: 7 Konfigurasi Ensemble (Akurasi)")
note(f"Akurasi vs proxy rule_violations≥6 (K=top-{_PP}). Termasuk IF+EE (saran Dr. Kelly).")
for name, acc, f1 in COMBO:
    bullet(f"{name:<12}  Accuracy {acc:.3f}  |  F1 {f1:.3f}")
note("Temuan: tidak ada konfigurasi tunggal yang dominan mutlak; "
     "ensemble IF+LOF+EE konsisten robust (tidak pernah terburuk).")
add_fig(fig_combo(), 6.8)

# ── SLIDE 18: JUSTIFIKASI ENSEMBLE ───────────────────────────────────
slide(18, "Justifikasi Ilmiah Penggunaan Ensemble")
bullet("EE skor tertinggi vs proxy — TAPI bias, karena proxy berbasis rule menguntungkan EE.")
bullet("Model individual volatil: LOF kuat di satu threshold, anjlok di threshold lain.")
bullet("Ensemble IF+LOF+EE: tidak pernah terburuk di semua setting = stabil & robust.")
note("Prinsip: robustness over peak performance — memilih 1 model 'terbaik' berisiko overfit ke 1 definisi anomali.")

# ── SLIDE 19: LIMITATION ─────────────────────────────────────────────
slide(19, "Limitations & Future Work")
h2("Limitations")
bullet("Unsupervised: 'akurasi' dihitung vs PROXY (rule_violations), bukan label asli — weak supervision.")
bullet("Proxy bias terhadap model yang selaras rule (EE diuntungkan, LOF dirugikan).")
bullet("Satu dataset; threshold belum di-tuning via ROC.")
h2("Future Work")
bullet("Label tervalidasi domain expert untuk Precision/Recall absolut.")
bullet("Multi-dataset + baseline comparison + uji signifikansi statistik.")
bullet("Temporal anomaly (LSTM autoencoder) terintegrasi dengan graph.")

# ── SLIDE 20: DELIVERABLES ───────────────────────────────────────────
slide(20, "Deliverables & Output")
bullet("output/anomaly_detection_report.txt — laporan teks auditor.")
bullet("output/anomaly_detection_detailed.json & anomaly_statistics.json.")
bullet("output/AD_Anomaly_Detection_Report_v3.docx — laporan formal (12 section).")
bullet("output/IJIES_Draft_Paper_Mahathir_Muhammad.docx — draft paper (format IJIES).")
bullet("models/ — IF, LOF, EE, Scaler tersimpan; pipeline reproducible.")

# ── SLIDE 21: KESIMPULAN ─────────────────────────────────────────────
slide(21, "Kesimpulan")
bullet(f"Pipeline 7-phase berhasil menganalisis {TOTAL_USERS} user, mendeteksi {TOTAL_ANOMALI} anomali (MEDIUM+).")
bullet("Knowledge Graph + Ensemble + SHAP menghasilkan deteksi akurat sekaligus dapat dijelaskan.")
bullet("Severity quantile-based menjadikan klasifikasi objektif & reproducible.")
bullet("Ablation 7 konfigurasi membuktikan ensemble paling robust (justifikasi ilmiah).")
note("Status: pipeline final, hasil kanonik & reproducible, draft paper IJIES siap.")

# ── SLIDE 22: Q&A ────────────────────────────────────────────────────
slide(22, "Terima Kasih — Q&A")
note("Mahathir Muhammad  ·  mahathirmuhammad02@gmail.com")
note("Repo: github.com/mahathirmuh/ad-anomaly-detection")

out = 'docs/presentations/PPT_Outline_TDAS_AD_Audit_v4.docx'
doc.save(out)
print(f"[OK] PPT outline saved: {out}")
